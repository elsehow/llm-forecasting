"""Generate Forecast Audit for Kapoor & Narayanan (2023) — Leakage and the Reproducibility Crisis."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elm)


def add_styled_paragraph(doc, text, style="Normal", bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def build_document():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ========================================================
    # TITLE PAGE
    # ========================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FORECAST AUDIT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        '"Leakage and the Reproducibility Crisis\nin Machine-Learning-Based Science"'
    )
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    source_line = doc.add_paragraph()
    source_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source_line.add_run("Source: Kapoor & Narayanan, Patterns 4, September 2023")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run("Audit date: February 25, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line = doc.add_paragraph()
    prep_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line.add_run("Prepared for: Ann Cleveland")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line2 = doc.add_paragraph()
    prep_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line2.add_run("Prepared by: Forecasting Research Institute")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This audit evaluates the key assumptions underlying Kapoor & Narayanan\u2019s "
        "2023 paper on data leakage and the reproducibility crisis in ML-based science. "
        "For each assumption, we extract the report\u2019s implicit claim, assess the "
        "evidence provided, and assign a calibrated probability using structured "
        "forecasting methods."
    )

    doc.add_paragraph(
        "The paper\u2019s headline conclusion is that data leakage is a widespread, "
        "systemic problem that inflates ML performance claims, and that when corrected, "
        "complex ML models lose their advantages over simpler traditional methods. "
        "Model info sheets are proposed as the primary intervention. Our audit finds "
        "the core empirical finding is well-supported, but the proposed solutions face "
        "significant adoption barriers."
    )

    # Summary table
    doc.add_paragraph("")
    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Assumption", "Sensitivity", "P(Report is Right)", "Verdict"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A1A2E")

    rows_data = [
        ["Leakage correction eliminates ML advantage", "HIGH", "72%", "Probably right"],
        ["Model info sheets will be adopted & effective", "HIGH", "8%", "Likely wrong"],
        ["The crisis is systemic across fields", "MEDIUM", "75%", "Probably right"],
        ["The problem is worsening", "MEDIUM", "60%", "Plausible but uncertain"],
        ["Peer review is insufficient to catch leakage", "HIGH", "92%", "Probably right"],
        ["Interdisciplinary teams will resolve the gap", "MEDIUM", "30%", "Uncertain"],
    ]

    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if col_idx == 1:  # Sensitivity column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if val == "HIGH":
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    run.bold = True
            elif col_idx == 2:  # Probability column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
        # Alternate row shading
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F5F5FA")

    doc.add_paragraph("")

    # Key findings
    doc.add_heading("Key Findings", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Where the report is strongest: ")
    run.bold = True
    p.add_run(
        "The core empirical finding\u2014that data leakage inflates ML performance "
        "claims and that correcting it eliminates ML\u2019s advantage over traditional "
        "models\u2014is well-supported (72%) and increasingly confirmed across fields "
        "including bioinformatics, neuroimaging, and clinical prediction. The diagnosis "
        "of peer review\u2019s inability to catch leakage (92%) is also strongly supported "
        "by the near-zero adoption of leakage-specific review requirements after nearly "
        "3 years of advocacy."
    )

    p = doc.add_paragraph()
    run = p.add_run("Where the report over-reaches: ")
    run.bold = True
    p.add_run(
        "The proposed solution\u2014model info sheets\u2014faces severe adoption barriers. "
        "Based on historical precedent (TRIPOD took 7+ years for even 20% journal "
        "endorsement), mandatory leakage-specific reporting requirements at 3+ major "
        "journals is only 8% likely by 2027. The recommendation that interdisciplinary "
        "teams will resolve the expertise gap (30%) has no empirical support and faces "
        "its own well-documented coordination challenges."
    )

    p = doc.add_paragraph()
    run = p.add_run("The bottom line: ")
    run.bold = True
    p.add_run(
        "The paper\u2019s diagnosis is substantially correct\u2014leakage is widespread, "
        "consequential, and under-addressed. But the prognosis is worse than the paper "
        "suggests: the proposed interventions are unlikely to be adopted at scale, and "
        "the institutional mechanisms needed to fix the problem (journal policy change, "
        "reviewer expertise, interdisciplinary collaboration) all face their own deep "
        "structural barriers. Decision-makers should treat ML-based scientific claims "
        "with skepticism, particularly in fields that rely on tabular data, but should "
        "not expect the problem to be solved by voluntary adoption of reporting standards."
    )

    doc.add_page_break()

    # ========================================================
    # HOW TO READ THIS AUDIT
    # ========================================================
    doc.add_heading("How to Read This Audit", level=1)

    doc.add_paragraph(
        "Every analytical report rests on assumptions\u2014claims the author treats as true "
        "without fully defending them. Some are well-supported. Others are not. "
        "When a decision-maker relies on a report, they inherit its assumptions."
    )

    doc.add_paragraph(
        "This audit makes those assumptions explicit. For each one, we ask:"
    )

    bullets = [
        "What does the report assume?",
        "What evidence does it provide?",
        "How much would the conclusion change if this assumption is wrong? (Sensitivity)",
        "How likely is the assumption to be correct? (Probability)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Probabilities are generated using structured forecasting: base rate construction, "
        "evidence weighting, decomposition analysis, and calibration checks. They represent "
        "our best estimate given available evidence, not certainty."
    )

    p = doc.add_paragraph()
    run = p.add_run('The "Given that" framing: ')
    run.bold = True
    p.add_run(
        "For each assumption, we describe the sensitivity using presuppositional framing: "
        "\"Given that this assumption is wrong, here is how the conclusion changes.\" "
        "Research shows this framing produces more accurate conditional reasoning than "
        "hypothetical framing (\"If this were wrong...\")."
    )

    doc.add_page_break()

    # ========================================================
    # DETAILED ASSUMPTION ANALYSIS
    # ========================================================
    doc.add_heading("Detailed Assumption Analysis", level=1)

    # --- ASSUMPTION 1 ---
    doc.add_heading("Assumption 1: Correcting Leakage Eliminates ML\u2019s Advantage Over Traditional Models", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "When leakage is corrected, complex ML models (Random Forests, Adaboost, GBT) "
        "no longer outperform decades-old logistic regression in civil war prediction. "
        "The implication is that this generalizes to other applied science fields."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Detailed replication of 4 papers in civil war prediction, all published in top "
        "political science journals. Each suffered from different forms of leakage. "
        "When corrected, ML models performed no better than baselines. The paper also "
        "cites similar findings in predicting children\u2019s life outcomes and recidivism, "
        "though these were not framed as leakage corrections."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the civil war prediction case study does not generalize\u2014that is, "
        "complex ML models genuinely outperform traditional methods in most applied fields "
        "even after correcting for leakage\u2014the paper\u2019s central narrative collapses. "
        "It becomes a niche finding about one subfield rather than evidence of a systemic crisis. "
        "The urgency of the proposed interventions diminishes dramatically."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (72%): ")
    run.bold = True
    p.add_run(
        "We estimate a 72% probability that systematic replication studies in fields "
        "outside political science will confirm this finding by 2027. At least one "
        "published paper\u2014Bernett et al. (2024) in Briefings in Bioinformatics, which "
        "found that deep learning models for protein-protein interaction prediction "
        "performed at random once leakage was corrected\u2014already comes close to meeting "
        "this criterion. Active research pipelines in neuroimaging, clinical psychiatry, "
        "and clinical prediction are converging on similar conclusions. The broader finding "
        "that ML does not consistently outperform logistic regression on tabular/structured "
        "data is being established through multiple meta-analyses. The main caveat: in "
        "domains with genuinely high-dimensional data (computer vision, NLP), ML\u2019s "
        "advantage is real and not an artifact of leakage."
    )

    # --- ASSUMPTION 2 ---
    doc.add_heading("Assumption 2: Model Info Sheets Will Be Adopted and Reduce Leakage", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "A 21-question model info sheet template will help researchers detect and prevent "
        "leakage. Journals should encourage or require their use. This is positioned as "
        "the paper\u2019s primary practical recommendation."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The template is presented and explained. No empirical evidence of adoption or "
        "effectiveness is provided (the paper was proposing a new tool, not evaluating one). "
        "The authors acknowledge limitations: info sheets can\u2019t be verified without "
        "computational reproducibility, they require ML expertise, and incorrect claims "
        "might provide false assurance."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that model info sheets are not widely adopted, the paper\u2019s proposed "
        "solution fails. The diagnosis stands (leakage is widespread) but the prognosis "
        "changes: the field lacks a viable intervention. This shifts the policy response "
        "from \u201cadopt these tools\u201d toward more structural approaches\u2014mandatory "
        "code review, automated leakage detection, or fundamental changes to how ML-based "
        "scientific claims are evaluated."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (8%): ")
    run.bold = True
    p.add_run(
        "We estimate only an 8% probability that 3+ major journals will adopt mandatory "
        "or strongly recommended leakage-specific reporting requirements by 2027. The "
        "strongest reference class is TRIPOD, a prediction model reporting checklist that "
        "took 7 years to reach 20% endorsement among high-impact journals, with just 5% "
        "making it mandatory. Model info sheets face additional headwinds: they target a "
        "narrower problem, lack a centralized adoption body (like the ICMJE that pushed "
        "CONSORT), and must compete with broader ML checklists like REFORMS. Nature\u2019s "
        "Machine Learning Checklist v1.1 already includes a leakage item, but formal "
        "policy adoption with enforcement across 3+ journals in 22 months is historically "
        "unprecedented for this type of intervention."
    )

    # --- ASSUMPTION 3 ---
    doc.add_heading("Assumption 3: The Crisis Is Systemic Across Scientific Fields", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "Based on a survey of 22 prior reviews across 17 fields, collectively identifying "
        "294 affected papers, leakage is a systemic crisis in ML-based science\u2014not "
        "isolated to a few fields or a handful of careless researchers."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "A comprehensive cross-disciplinary survey table (Figure 1) mapping leakage types "
        "by field. The authors note this is a \u201clower bound\u201d since they only "
        "examined paper content, not code. Some fields had only 1 affected paper in the survey. "
        "The methodology relied on existing reviews that had already identified problems, "
        "creating potential selection bias toward fields with active methodological debate."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that the crisis is concentrated in a few fields rather than being truly "
        "systemic, the policy response narrows. Instead of field-agnostic solutions (like "
        "model info sheets for all ML-based science), targeted interventions for the most "
        "affected fields (medicine, neuroimaging, political science) would be more appropriate. "
        "The urgency of cross-disciplinary coordination diminishes."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our assessment (75%): ")
    run.bold = True
    p.add_run(
        "We estimate a 75% probability that the crisis is genuinely systemic as described. "
        "The original survey has since been updated to 648 papers across 30 fields (May 2024), "
        "and field-specific reviews continue to emerge in neuroimaging, bioinformatics, "
        "genomics, clinical prediction, and psychiatry. However, independent cross-disciplinary "
        "confirmation is unlikely (4% that 2+ independent surveys will be published by 2027), "
        "largely because such surveys are labor-intensive and Kapoor & Narayanan are occupying "
        "the niche. The claim rests primarily on their own expanding survey plus converging "
        "field-specific evidence."
    )

    # --- ASSUMPTION 4 ---
    doc.add_heading("Assumption 4: The Problem Is Worsening Due to Feedback Loops", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "ML adoption is accelerating in applied sciences, non-replicable papers are cited "
        "more than replicable ones, and this creates a feedback loop of overoptimism. "
        "The problem is getting worse, not better."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Figure 2 shows a sharp increase in civil war prediction papers using ML from "
        "2013\u20132022. The paper cites Serra-Garcia & Gneezy (2021) showing non-replicable "
        "papers are cited more. The \u201claissez-faire approach\u201d of independent "
        "communities discovering the same pitfalls is presented as evidence of spreading errors."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that the problem is stabilizing or improving\u2014due to growing awareness, "
        "new checklists, and better education\u2014the urgency of immediate intervention "
        "decreases. The paper\u2019s crisis framing would be overstated, and a more gradual "
        "approach to reform would be justified."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our assessment (60%): ")
    run.bold = True
    p.add_run(
        "We estimate a 60% probability that the problem is indeed getting worse in absolute "
        "terms. ML adoption is unquestionably accelerating\u2014more fields, more papers, "
        "more researchers using off-the-shelf tools. The citation bias toward positive results "
        "is well-documented. However, awareness has grown substantially since 2023: the "
        "REFORMS checklist, conference reproducibility requirements, and emerging automated "
        "detection tools provide countervailing forces. The rate of formal corrections "
        "remains very low (only 12% likely to double by 2027), suggesting the correction "
        "mechanism is largely broken regardless of whether the underlying problem worsens. "
        "The net picture: more leaky papers being produced, but also more being identified\u2014"
        "a race between production and detection."
    )

    # --- ASSUMPTION 5 ---
    doc.add_heading("Assumption 5: Existing Peer Review Cannot Catch Leakage", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "Current peer review practices are insufficient to detect leakage. Reviewers "
        "don\u2019t examine code, many fields lack ML expertise among reviewers, and "
        "existing reporting standards don\u2019t address leakage. Even when code is "
        "available, reviewing it is time-consuming and requires specialized knowledge."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "All 294 leaky papers passed peer review at established journals, including "
        "4 papers published in top-10 political science and international relations "
        "journals. The paper notes that current peer review doesn\u2019t require code "
        "disclosure and that checklist adoption is slow."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that peer review can actually catch leakage effectively, the proposed "
        "interventions (model info sheets, new checklists) are largely unnecessary\u2014"
        "the existing system just needs better training and resources. The urgency of "
        "developing new tools diminishes, and the policy response shifts to reviewer "
        "education and resource allocation."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (92%): ")
    run.bold = True
    p.add_run(
        "We estimate a 92% probability that peer review remains insufficient to catch "
        "leakage. Our forecast that 5+ journals will implement mandatory leakage-specific "
        "code review by 2027 came in at just 3%. Zero journals currently meet this "
        "criterion. Even the most advanced code-review-implementing journals conduct "
        "general reviews without leakage-specific instructions. The REFORMS checklist, "
        "published nearly 2 years ago, has seen no evidence of mandatory journal adoption. "
        "A severe peer review crisis\u2014reviewers declining at record rates, median "
        "review times stretching\u2014makes adding new reviewer burdens especially unlikely. "
        "This is the paper\u2019s strongest supported assumption."
    )

    # --- ASSUMPTION 6 ---
    doc.add_heading("Assumption 6: Interdisciplinary Collaboration Will Resolve the Expertise Gap", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "Recommendation R3 calls for interdisciplinary teams combining ML and domain "
        "expertise. The paper argues the expertise gap\u2014domain scientists lack ML "
        "knowledge, ML researchers lack domain knowledge\u2014is a key driver of leakage."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "None. This is presented as a recommendation without empirical support. The paper "
        "does not test whether interdisciplinary teams produce fewer leakage errors, "
        "nor does it cite evidence that they do."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that interdisciplinary collaboration does not substantially reduce leakage "
        "rates\u2014because leakage can be subtle enough to evade even ML experts, or "
        "because collaboration introduces its own coordination costs\u2014the field needs "
        "structural solutions (automated tools, mandatory code verification) rather than "
        "relying on team composition changes."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our assessment (30%): ")
    run.bold = True
    p.add_run(
        "We estimate a 30% probability that interdisciplinary collaboration is an effective "
        "solution to the leakage problem. No empirical evidence exists that team composition "
        "affects leakage rates. Only a 4% chance that such evidence will be published by "
        "2027, given methodological barriers: operationalizing \u201cinterdisciplinary\u201d "
        "is contested, the required study design (leakage audit + team composition coding "
        "across hundreds of papers) is labor-intensive, and confounding is severe (well-"
        "resourced labs are both more interdisciplinary and better-resourced for methodological "
        "rigor). The recommendation is intuitively appealing but empirically vacant."
    )

    doc.add_page_break()

    # ========================================================
    # IMPLICATIONS FOR DECISION-MAKERS
    # ========================================================
    doc.add_heading("Implications for Decision-Makers", level=1)

    doc.add_heading("What to act on now", level=2)

    bullets_act = [
        (
            "Discount ML performance claims in applied science (72% the finding generalizes). ",
            "When ML papers in fields like medicine, social science, or environmental science "
            "claim dramatic performance gains over traditional methods, treat these with "
            "skepticism\u2014especially if the paper does not explicitly address data leakage. "
            "This is particularly relevant for policy-relevant prediction tasks (recidivism, "
            "health outcomes, conflict prediction) where inflated claims can drive real-world "
            "decisions."
        ),
        (
            "Do not rely on peer review to catch leakage (92% it cannot). ",
            "The fact that a paper passed peer review at a top journal provides essentially "
            "no assurance about leakage. Decision-makers who rely on ML-based scientific "
            "findings should commission independent code review or require computational "
            "reproducibility before acting on high-stakes claims."
        ),
        (
            "Invest in automated leakage detection, not voluntary checklists. ",
            "Model info sheets and reporting standards are unlikely to achieve meaningful "
            "adoption (8%). Automated tools\u2014static analysis of ML code, LLM-based code "
            "review\u2014offer a more scalable path. Emerging tools like LeakageDetector "
            "and LLM-based leakage detection show promise, though they are not yet mature "
            "enough for deployment."
        ),
    ]

    for bold_text, normal_text in bullets_act:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_heading("What to watch", level=2)

    bullets_watch = [
        (
            "Field-specific replication studies. ",
            "The next 2 years will likely produce leakage-correction studies in neuroimaging, "
            "bioinformatics, and clinical prediction. If these consistently confirm that ML "
            "advantages disappear after correction, the case for systemic skepticism strengthens "
            "further. If they find ML advantages survive in some fields, the picture becomes "
            "more nuanced."
        ),
        (
            "Journal policy changes. ",
            "Nature\u2019s ML Checklist and the REFORMS initiative represent the leading "
            "edge of institutional response. If major publishers adopt mandatory leakage "
            "reporting in the next 2 years, the outlook improves. If they don\u2019t (which "
            "is more likely), the correction mechanism remains broken."
        ),
        (
            "Automated detection maturity. ",
            "Tools that can scan ML code for common leakage patterns could transform "
            "enforcement. Monitor whether automated leakage detection achieves journal-grade "
            "reliability by 2027. If so, the structural problem becomes solvable without "
            "relying on human reviewer expertise."
        ),
    ]

    for bold_text, normal_text in bullets_watch:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_page_break()

    # ========================================================
    # METHODOLOGY NOTE
    # ========================================================
    doc.add_heading("Methodology", level=1)

    doc.add_paragraph(
        "Each assumption was converted into a concrete, time-bound forecasting question "
        "and evaluated using structured forecasting methods:"
    )

    method_steps = [
        "Base rate construction from historical reference classes",
        "Evidence gathering from domain research, prediction markets, and recent publications",
        "Bayesian updating based on specific factors for and against",
        "Decomposition analysis (breaking the question into independent paths to resolution)",
        "Calibration checks and stress testing against sensitivity ranges",
    ]
    for step in method_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph(
        "Forecasts represent calibrated probabilities\u2014our best estimate of likelihood "
        "given available evidence. A 60% forecast means we expect the event to occur in "
        "roughly 6 out of 10 similar situations. These are not expressions of confidence "
        "in our analysis; they are estimates of real-world probability."
    )

    doc.add_paragraph(
        "Sensitivity framing uses presuppositional (\"Given that X\") rather than "
        "hypothetical (\"If X were true\") language. Research on conditional reasoning "
        "shows that presuppositional framing produces more accurate counterfactual "
        "analysis by treating the alternative scenario as a concrete state of the world "
        "rather than an abstract possibility."
    )

    doc.add_paragraph("")

    doc.add_heading("Forecasting Questions", level=2)

    doc.add_paragraph(
        "The following specific forecasting questions were used to operationalize "
        "each assumption:"
    )

    forecast_qs = [
        (
            "Assumption 1: ",
            "Will a systematic replication study correcting for data leakage in an applied ML field "
            "outside political science, published in a peer-reviewed journal by Dec 2027, find that "
            "complex ML models no longer significantly outperform traditional statistical baselines? "
            "Direct forecast: 72%."
        ),
        (
            "Assumption 2: ",
            "Will at least 3 major journals adopt mandatory or strongly recommended leakage-specific "
            "reporting requirements by Dec 2027? Direct forecast: 8%."
        ),
        (
            "Assumption 3: ",
            "Will at least 2 independent cross-disciplinary surveys of ML leakage be published by "
            "Dec 2027, each identifying 100+ affected papers? Direct forecast: 4%. "
            "Assessment adjusted upward to 75% to reflect that the underlying claim can be "
            "supported by converging field-specific evidence without formal cross-disciplinary surveys."
        ),
        (
            "Assumption 4: ",
            "Will the annual rate of corrections/retractions citing data leakage increase by 100% "
            "from 2023 to 2027? Direct forecast: 12%. "
            "Assessment adjusted upward to 60% because the problem can worsen (more leaky papers "
            "being published) even while the correction mechanism remains broken."
        ),
        (
            "Assumption 5: ",
            "Will at least 5 journals implement mandatory leakage-specific code review by Dec 2027? "
            "Direct forecast: 3%. Assessment inverted: the 97% probability that reform does NOT "
            "happen supports the report\u2019s claim that peer review is insufficient."
        ),
        (
            "Assumption 6: ",
            "Will a published empirical study demonstrate that interdisciplinary teams have "
            "significantly lower leakage rates by Dec 2027? Direct forecast: 4%. "
            "Assessment adjusted upward to 30% to reflect that the recommendation could be "
            "partially correct even without published evidence."
        ),
    ]

    for bold_text, normal_text in forecast_qs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    # ========================================================
    # SAVE
    # ========================================================
    output_path = "/Users/elsehow/Downloads/audit/ann/FRI_Forecast_Audit_Kapoor_Narayanan_Leakage.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
