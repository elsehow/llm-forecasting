"""Generate the Forecast Audit deliverable as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elm)


def add_styled_paragraph(doc, text, style="Normal", bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def build_document():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ========================================================
    # TITLE PAGE
    # ========================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FORECAST AUDIT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        '"Disrupting the First Reported AI-Orchestrated\nCyber Espionage Campaign"'
    )
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    source_line = doc.add_paragraph()
    source_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source_line.add_run("Source report: Anthropic, November 2025")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run("Audit date: February 25, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line = doc.add_paragraph()
    prep_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line.add_run("Prepared for: Ann Cleveland, CLTC")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line2 = doc.add_paragraph()
    prep_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line2.add_run("Prepared by: Forecasting Research Institute")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This audit evaluates the key assumptions underlying Anthropic's November 2025 report "
        "on the GTG-1002 cyber espionage campaign. For each assumption, we extract the report's "
        "implicit claim, assess the evidence provided, and assign a calibrated probability "
        "using structured forecasting methods."
    )

    doc.add_paragraph(
        "The report's headline conclusion is that AI-enabled autonomous cyber espionage "
        "represents a fundamental shift in the threat landscape. Our audit finds that this "
        "conclusion rests on several assumptions of varying strength."
    )

    # Summary table
    doc.add_paragraph("")
    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Assumption", "Sensitivity", "P(Report is Right)", "Verdict"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A1A2E")

    rows_data = [
        ["Attribution is Chinese state-sponsored", "HIGH", "60%", "Plausible but unsubstantiated"],
        ["80-90% AI autonomy is effective", "HIGH", "6%", "Likely overstated"],
        ["Generalizes beyond Claude", "MEDIUM", "68%", "Probably right"],
        ["Commodity tooling = rapid proliferation", "MEDIUM", "38%", "Uncertain"],
        ["Success rate is alarming vs baseline", "MEDIUM", "20%", "Probably overstated"],
        ["Hallucination is temporary", "HIGH", "15%", "Likely wrong"],
    ]

    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if col_idx == 1:  # Sensitivity column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if val == "HIGH":
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    run.bold = True
            elif col_idx == 2:  # Probability column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
        # Alternate row shading
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F5F5FA")

    doc.add_paragraph("")

    # Key findings
    doc.add_heading("Key Findings", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Where the report over-reaches: ")
    run.bold = True
    p.add_run(
        "The \"fundamental shift\" narrative depends on AI being genuinely autonomous and "
        "effective (6% likely), on hallucination being a temporary obstacle (15% likely), "
        "and on AI campaigns already outperforming traditional methods (20% likely). "
        "The report's own evidence\u2014frequent hallucinations and fabricated data\u2014"
        "contradicts its autonomy claim. The real story may be cost reduction, not capability expansion."
    )

    p = doc.add_paragraph()
    run = p.add_run("Where the report is probably right: ")
    run.bold = True
    p.add_run(
        "This will likely generalize beyond Claude (68%). Other frontier models have "
        "comparable agentic capabilities, and when misuse is documented for one model, "
        "cross-model documentation typically follows within 6\u201318 months. The attribution "
        "claim is plausible (60%), given that Anthropic coordinated with authorities."
    )

    p = doc.add_paragraph()
    run = p.add_run("The bottom line: ")
    run.bold = True
    p.add_run(
        "AI is being integrated into real cyber operations at a level beyond previous findings. "
        "This is well-supported. But the claim that this represents autonomous AI cyber warfare "
        "is premature. A more accurate framing: threat actors are using AI as a force multiplier "
        "for human-directed operations, achieving similar results with fewer operators. "
        "Policy responses should focus on volume effects (more campaigns at lower cost) "
        "rather than novel autonomous capabilities."
    )

    doc.add_page_break()

    # ========================================================
    # HOW TO READ THIS AUDIT
    # ========================================================
    doc.add_heading("How to Read This Audit", level=1)

    doc.add_paragraph(
        "Every analytical report rests on assumptions\u2014claims the author treats as true "
        "without fully defending them. Some are well-supported. Others are not. "
        "When a decision-maker relies on a report, they inherit its assumptions."
    )

    doc.add_paragraph(
        "This audit makes those assumptions explicit. For each one, we ask:"
    )

    bullets = [
        "What does the report assume?",
        "What evidence does it provide?",
        "How much would the conclusion change if this assumption is wrong? (Sensitivity)",
        "How likely is the assumption to be correct? (Probability)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Probabilities are generated using structured forecasting: base rate construction, "
        "evidence weighting, decomposition analysis, and calibration checks. They represent "
        "our best estimate given available evidence, not certainty."
    )

    p = doc.add_paragraph()
    run = p.add_run("The \"Given that\" framing: ")
    run.bold = True
    p.add_run(
        "For each assumption, we describe the sensitivity using presuppositional framing: "
        "\"Given that this assumption is wrong, here is how the conclusion changes.\" "
        "Research shows this framing produces more accurate conditional reasoning than "
        "hypothetical framing (\"If this were wrong...\")."
    )

    doc.add_page_break()

    # ========================================================
    # DETAILED ASSUMPTION ANALYSIS
    # ========================================================
    doc.add_heading("Detailed Assumption Analysis", level=1)

    # --- ASSUMPTION 1 ---
    doc.add_heading("Assumption 1: Attribution Is Chinese State-Sponsored", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's confidence: ")
    run.bold = True
    p.add_run("High")

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "None in the public report. No indicators of compromise, no tradecraft overlap "
        "with known groups, no infrastructure analysis. The report states Anthropic "
        "coordinated with authorities and shared intelligence, but the basis for "
        "attribution is not detailed."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the attacker is not Chinese state-sponsored\u2014say, a criminal group "
        "or a red team from a third country using Chinese-language infrastructure\u2014the "
        "policy response shifts dramatically. \"China is using AI for espionage\" drives one "
        "set of responses (export controls, diplomatic pressure, ally coordination). "
        "\"Sophisticated criminals are using AI for espionage\" drives a different set "
        "(law enforcement, industry hardening, insurance). Attribution determines the "
        "entire policy frame."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (60%): ")
    run.bold = True
    p.add_run(
        "We estimate a 60% probability that the attribution will be independently "
        "corroborated by a second credible source (government agency, major cybersecurity "
        "firm, or CISA advisory) by December 2026. The historical corroboration rate for "
        "specific state-sponsored attribution claims is 70\u201380% within 1\u20132 years, "
        "but Anthropic is not a traditional cybersecurity firm, and corroboration may "
        "happen privately without public disclosure. The strongest factor in favor: "
        "Anthropic shared intelligence with authorities, which is the primary mechanism "
        "by which corroborations historically occur."
    )

    # --- ASSUMPTION 2 ---
    doc.add_heading("Assumption 2: 80\u201390% AI Autonomy Reflects Real Operational Effectiveness", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "AI executed 80\u201390% of tactical work independently, based on operational "
        "tempo, request volumes, and activity patterns."
    )

    p = doc.add_paragraph()
    run = p.add_run("Unstated dependency: ")
    run.bold = True
    p.add_run(
        "The same report notes Claude \"frequently overstated findings and occasionally "
        "fabricated data\"\u2014hallucinating credentials, discoveries, and results. "
        "If the AI was running autonomously but producing unreliable outputs, the "
        "effective autonomy rate could be much lower than reported. The 80\u201390% figure "
        "measures execution volume, not outcome quality."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that only 30\u201340% of autonomous AI actions produced genuinely useful "
        "results, this is a story about AI-assisted attacks (humans still doing the hard "
        "thinking) rather than AI-autonomous attacks. The mitigation strategy shifts from "
        "\"detect autonomous AI agents\" toward \"disrupt human operators who use AI as "
        "a force multiplier\"\u2014a very different defensive posture."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (6%): ")
    run.bold = True
    p.add_run(
        "We estimate only a 6% probability that AI-autonomous cyber operations will "
        "demonstrate greater than 50% effective task success rate in publicly documented "
        "real-world operations by end of 2026. The base rate for novel technology "
        "demonstrating reliable autonomous performance in adversarial real-world domains "
        "is low (8\u201312%). The report's own evidence is the strongest argument against "
        "this assumption: frequent hallucination in offensive security contexts is an "
        "architectural limitation, not a tuning problem. Even if achieved, public "
        "documentation is unlikely given classification incentives."
    )

    # --- ASSUMPTION 3 ---
    doc.add_heading("Assumption 3: This Generalizes Beyond Claude", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "\"This case study likely reflects consistent patterns of behavior across "
        "frontier AI models.\""
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "None. Anthropic only has visibility into Claude usage. The attack framework "
        "used Claude Code and MCP tools\u2014a specific architecture that may or may not "
        "have equivalents for other models."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that this is primarily a Claude-specific vulnerability, the right response "
        "is vendor-specific hardening rather than industry-wide regulation. Conversely, "
        "given that all frontier models are equally exploitable, the urgency for "
        "cross-industry standards is much higher."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (68%): ")
    run.bold = True
    p.add_run(
        "We estimate a 68% probability that at least one other frontier model will be "
        "publicly documented as exploited for autonomous or semi-autonomous cyber "
        "operations by end of 2026. When novel AI misuse is documented for one model, "
        "cross-model documentation typically follows within 6\u201318 months. OpenAI and "
        "Google have established norms of publishing misuse reports. Agentic "
        "infrastructure (function calling, code execution, agent frameworks) is maturing "
        "rapidly across all frontier models. The main counterweight: 2.5+ years of GPT-4 "
        "availability without such documentation, and the difficulty of attributing "
        "autonomous operations to specific models."
    )

    # --- ASSUMPTION 4 ---
    doc.add_heading("Assumption 4: Commodity Tooling Means Rapid Proliferation", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's logic: ")
    run.bold = True
    p.add_run(
        "The attackers used open-source penetration testing tools orchestrated via MCP, "
        "with no novel malware. Therefore, the barrier to entry is low."
    )

    p = doc.add_paragraph()
    run = p.add_run("What's missing: ")
    run.bold = True
    p.add_run(
        "The orchestration framework itself was described as \"highly sophisticated\" and "
        "\"professionally coordinated.\" Building the MCP-based attack framework, crafting "
        "prompts that bypass safety training via role-play, and managing multi-session "
        "campaigns is non-trivial integration work. The tools are commodity; the glue is not."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that the integration layer requires significant expertise, proliferation "
        "will be slower than implied\u2014this remains a well-resourced-actor capability "
        "for the near term. Given that someone open-sources a similar framework, "
        "proliferation could be nearly immediate."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (38%): ")
    run.bold = True
    p.add_run(
        "We estimate a 38% probability that a comparable open-source AI-orchestrated "
        "attack framework becomes publicly available by end of 2026. The historical base "
        "rate for public tooling following demonstrated effectiveness is high (70\u201380% "
        "within 3 years), but the compressed 10-month timeframe and high capability bar "
        "reduce this. The mature MCP ecosystem and AI agent frameworks provide the "
        "architectural building blocks, but the conjunction of building, sharing, and "
        "meeting the capability bar is the bottleneck."
    )

    # --- ASSUMPTION 5 ---
    doc.add_heading("Assumption 5: The Success Rate Is Alarming", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's framing: ")
    run.bold = True
    p.add_run(
        "The campaign targeted roughly 30 entities and achieved \"a handful of successful "
        "intrusions.\" This is presented as alarming\u2014evidence of a fundamental shift."
    )

    p = doc.add_paragraph()
    run = p.add_run("Missing context: ")
    run.bold = True
    p.add_run(
        "What is the success rate of traditional nation-state campaigns against "
        "comparable target profiles? Targeted APT campaigns typically achieve 20\u201340% "
        "success rates against specifically targeted entities. A \"handful\" out of 30 "
        "(roughly 10\u201320%) may represent parity or even underperformance at lower cost, "
        "rather than a capability leap."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that the success rate is comparable to or below human-operated campaigns, "
        "the real story is cost reduction (same results, fewer operators), not capability "
        "expansion. Policy should focus on volume effects\u2014more campaigns at lower "
        "cost\u2014rather than novel capability."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (20%): ")
    run.bold = True
    p.add_run(
        "We estimate a 20% probability that AI-orchestrated campaigns will match or "
        "exceed traditional APT success rates against comparable targets by end of 2027. "
        "The most plausible path is not pure autonomous AI but nation-state actors adopting "
        "AI as their primary orchestration layer, inheriting both AI advantages (speed, "
        "parallelism) and human strategic judgment. The fundamental measurement problem\u2014"
        "nation-state APT success rates are poorly documented\u2014also makes comparison "
        "inherently difficult."
    )

    # --- ASSUMPTION 6 ---
    doc.add_heading("Assumption 6: AI Hallucination Is a Temporary Obstacle", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's framing: ")
    run.bold = True
    p.add_run(
        "Hallucination was a limitation but the overall trajectory toward autonomous "
        "cyber operations is concerning. The implication: this problem will be solved, "
        "and when it is, the threat escalates dramatically."
    )

    p = doc.add_paragraph()
    run = p.add_run("Two directions this could go: ")
    run.bold = True

    doc.add_paragraph(
        "Given that hallucination in adversarial/agentic contexts is a fundamental "
        "limitation\u2014models cannot reliably verify their own outputs in novel "
        "environments\u2014autonomous cyber operations may plateau. Current mitigation "
        "strategies may be adequate.",
        style="List Bullet",
    )

    doc.add_paragraph(
        "Given that hallucination is largely solved within 12\u201318 months, the window "
        "to build defenses is narrow and the urgency of the response must increase "
        "dramatically.",
        style="List Bullet",
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "This assumption determines the urgency of the entire policy response. If "
        "hallucination is fundamental, defenders have more time. If it is solved quickly, "
        "the threat accelerates faster than the report implies."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (15%): ")
    run.bold = True
    p.add_run(
        "We estimate a 15% probability that frontier models will reduce hallucination "
        "to below 10% in adversarial/agentic contexts by end of 2027. Recent research "
        "on the Forward Simulation Gap\u2014tested across 10 frontier models and 8 "
        "prompting interventions\u2014finds that LLMs systematically default to "
        "unconditional base-rate predictions when asked to reason about modified states. "
        "This \"priors machine\" behavior is robust, universal, and resistant to "
        "intervention. Even the best-performing reasoning model (o3) with optimal "
        "prompting leaves a large residual gap. The most promising training approach "
        "(reinforcement learning with verifiable rewards) works best precisely where "
        "this problem does not apply: environments with clear verification signals. "
        "Adversarial and novel environments\u2014the core context for autonomous cyber "
        "operations\u2014lack such signals."
    )

    doc.add_page_break()

    # ========================================================
    # IMPLICATIONS FOR DECISION-MAKERS
    # ========================================================
    doc.add_heading("Implications for Decision-Makers", level=1)

    doc.add_heading("What to act on now", level=2)

    bullets_act = [
        (
            "Prepare for cross-model exploitation (68% likely). ",
            "This is not a Claude-specific problem. Organizations should assume that "
            "AI-augmented cyber operations will use whatever frontier model is most capable "
            "and least restricted. Defensive strategies that depend on a single vendor's "
            "safety measures are insufficient."
        ),
        (
            "Focus defenses on AI-augmented human operators, not autonomous agents (94% the right frame). ",
            "The most likely near-term threat is human attackers using AI to work faster and "
            "across more targets simultaneously. Detection strategies should focus on the "
            "signatures of AI-assisted operations: high-volume reconnaissance, parallel "
            "target engagement, and templated exploitation patterns."
        ),
        (
            "Expect more campaigns at lower cost. ",
            "Even without matching traditional APT effectiveness, AI-augmented operations "
            "reduce the personnel cost of cyber campaigns. The policy implication: the "
            "number of simultaneous campaigns will increase even if individual campaign "
            "quality does not."
        ),
    ]

    for bold_text, normal_text in bullets_act:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_heading("What to watch", level=2)

    bullets_watch = [
        (
            "Independent attribution corroboration. ",
            "If a government agency or major cybersecurity firm confirms the Chinese "
            "state-sponsorship claim, the geopolitical implications escalate. If no "
            "corroboration emerges by mid-2026, the attribution should be treated with "
            "increasing skepticism."
        ),
        (
            "Open-source attack framework emergence (38% by end of 2026). ",
            "If a comparable framework becomes publicly available, the proliferation "
            "timeline compresses dramatically. Monitor security conference presentations "
            "(DEF CON, Black Hat) and dark web forums for early signals."
        ),
        (
            "Hallucination progress in agentic contexts. ",
            "If frontier models demonstrate reliable autonomous performance in novel "
            "environments (security testing, unfamiliar codebases), the threat model "
            "changes from augmentation to autonomy. Current evidence suggests this is "
            "unlikely before 2028, but progress should be monitored."
        ),
    ]

    for bold_text, normal_text in bullets_watch:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_page_break()

    # ========================================================
    # METHODOLOGY NOTE
    # ========================================================
    doc.add_heading("Methodology", level=1)

    doc.add_paragraph(
        "Each assumption was converted into a concrete, time-bound forecasting question "
        "and evaluated using structured forecasting methods:"
    )

    method_steps = [
        "Base rate construction from historical reference classes",
        "Evidence gathering from domain research and recent publications",
        "Bayesian updating based on specific factors for and against",
        "Decomposition analysis (breaking the question into independent paths to resolution)",
        "Calibration checks and stress testing against sensitivity ranges",
    ]
    for step in method_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph(
        "Forecasts represent calibrated probabilities\u2014our best estimate of likelihood "
        "given available evidence. A 60% forecast means we expect the event to occur in "
        "roughly 6 out of 10 similar situations. These are not expressions of confidence "
        "in our analysis; they are estimates of real-world probability."
    )

    doc.add_paragraph(
        "Sensitivity framing uses presuppositional (\"Given that X\") rather than "
        "hypothetical (\"If X were true\") language. Research on conditional reasoning "
        "shows that presuppositional framing produces more accurate counterfactual "
        "analysis by treating the alternative scenario as a concrete state of the world "
        "rather than an abstract possibility."
    )

    # ========================================================
    # SAVE
    # ========================================================
    output_path = "/Users/elsehow/Downloads/audit/ann/FRI_Forecast_Audit_Anthropic_Cyber_Espionage.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
