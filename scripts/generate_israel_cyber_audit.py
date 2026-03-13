"""Generate Forecast Audit deliverable for Israel National Cybersecurity Strategy 2025."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex},
    )
    shading.append(shading_elm)


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ========================================================
    # TITLE PAGE
    # ========================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FORECAST AUDIT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Israel National Cyber Security Strategy 2025"
    )
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    source_line = doc.add_paragraph()
    source_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source_line.add_run(
        "Source report: Israel National Cyber Directorate (INCD), February 2025"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run("Audit date: February 25, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line = doc.add_paragraph()
    prep_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line.add_run("Prepared for: Javier, Repsol")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line2 = doc.add_paragraph()
    prep_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line2.add_run("Prepared by: Forecasting Research Institute")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This audit evaluates the key assumptions underlying the Israel National "
        "Cyber Directorate\u2019s 2025 cybersecurity strategy. The strategy proposes "
        "a comprehensive national cyber defense built around the \u201cCyber Dome\u201d "
        "AI system, whole-of-society joint defense, workforce expansion, and "
        "international partnerships, with a 2028 implementation deadline. For each "
        "assumption, we extract the report\u2019s implicit claim, assess the evidence "
        "provided, and assign a calibrated probability using structured forecasting methods."
    )

    doc.add_paragraph(
        "The strategy\u2019s headline conclusion is that Israel can maintain its "
        "position as a global cybersecurity leader through integrated national defense, "
        "technological innovation, and international cooperation. Our audit finds a "
        "mixed picture: the strategy\u2019s threat assessment and international "
        "positioning are well-grounded, but its implementation timeline and workforce "
        "assumptions are significantly overstated."
    )

    # Summary table
    doc.add_paragraph("")
    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Assumption", "Sensitivity", "P(Report is Right)", "Verdict"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A1A2E")

    rows_data = [
        ["Cyber Dome delivers effective defense", "HIGH", "55%", "Plausible but uncertain"],
        ["Cyber workforce can scale fast enough", "HIGH", "13%", "Probably overstated"],
        ["Joint defense overcomes fragmentation", "HIGH", "25%", "Uncertain"],
        ["International partnerships remain robust", "MEDIUM", "95%", "Probably right"],
        ["Iran-centric threat model is adequate", "MEDIUM", "87%", "Probably right"],
        ["2028 implementation timeline achievable", "HIGH", "25%", "Uncertain"],
    ]

    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if col_idx == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if val == "HIGH":
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    run.bold = True
            elif col_idx == 2:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F5F5FA")

    doc.add_paragraph("")

    # Key findings
    doc.add_heading("Key Findings", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Where the strategy is well-grounded: ")
    run.bold = True
    p.add_run(
        "Israel\u2019s threat assessment is accurate. Iran and its proxies will "
        "almost certainly remain the dominant cyber adversary (87%), with operations "
        "intensifying rather than declining post-ceasefire. International partnerships "
        "are already expanding (95%)\u2014the Germany and India agreements signed in "
        "late 2025 and early 2026 demonstrate that cybersecurity cooperation is "
        "resilient to broader geopolitical headwinds."
    )

    p = doc.add_paragraph()
    run = p.add_run("Where the strategy over-reaches: ")
    run.bold = True
    p.add_run(
        "The 2028 implementation timeline (25%) and workforce scaling assumptions "
        "(13%) are the weakest links. Israel\u2019s tech workforce is currently "
        "shrinking for the first time in a decade, and the INCD has tried and failed "
        "to pass comprehensive cybersecurity legislation twice before. Elections "
        "scheduled for 2026 will reset the legislative calendar again. The education "
        "pipeline initiatives targeting elementary students will not produce workers "
        "within the strategy\u2019s timeframe."
    )

    p = doc.add_paragraph()
    run = p.add_run("The Cyber Dome question (55%): ")
    run.bold = True
    p.add_run(
        "The strategy\u2019s centerpiece is a genuine coin-flip. The system is in "
        "preliminary operation and Israel has strong incentives to demonstrate "
        "effectiveness, but no country has successfully deployed a nationwide AI "
        "cyber defense umbrella. The strongest argument in its favor is simply the "
        "high volume of attacks creating many opportunities for documented success."
    )

    p = doc.add_paragraph()
    run = p.add_run("The bottom line: ")
    run.bold = True
    p.add_run(
        "Israel\u2019s cybersecurity strategy is stronger on diagnosis than prescription. "
        "Its threat assessment and international positioning are sound. Its proposed "
        "solutions\u2014particularly the legislative reform, workforce expansion, and "
        "ambitious timeline\u2014face structural barriers that the strategy acknowledges "
        "but underestimates. Decision-makers should expect meaningful progress on "
        "Cyber Dome and international cooperation, but plan for a longer implementation "
        "horizon and persistent workforce constraints."
    )

    doc.add_page_break()

    # ========================================================
    # HOW TO READ THIS AUDIT
    # ========================================================
    doc.add_heading("How to Read This Audit", level=1)

    doc.add_paragraph(
        "Every analytical report rests on assumptions\u2014claims the author treats as true "
        "without fully defending them. Some are well-supported. Others are not. "
        "When a decision-maker relies on a report, they inherit its assumptions."
    )

    doc.add_paragraph(
        "This audit makes those assumptions explicit. For each one, we ask:"
    )

    for b in [
        "What does the report assume?",
        "What evidence does it provide?",
        "How much would the conclusion change if this assumption is wrong? (Sensitivity)",
        "How likely is the assumption to be correct? (Probability)",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Probabilities are generated using structured forecasting: base rate construction, "
        "evidence weighting, decomposition analysis, and calibration checks. They represent "
        "our best estimate given available evidence, not certainty."
    )

    p = doc.add_paragraph()
    run = p.add_run('The "Given that" framing: ')
    run.bold = True
    p.add_run(
        "For each assumption, we describe the sensitivity using presuppositional framing: "
        '"Given that this assumption is wrong, here is how the conclusion changes." '
        "Research shows this framing produces more accurate conditional reasoning than "
        'hypothetical framing ("If this were wrong...").'
    )

    doc.add_page_break()

    # ========================================================
    # DETAILED ASSUMPTION ANALYSIS
    # ========================================================
    doc.add_heading("Detailed Assumption Analysis", level=1)

    # --- ASSUMPTION 1: Cyber Dome ---
    doc.add_heading(
        "Assumption 1: The Cyber Dome Will Deliver Effective National Defense",
        level=2,
    )

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "The Cyber Dome\u2014an AI-powered, multi-layered defense system\u2014will "
        "provide holistic situational awareness, active threat defense, and "
        "early warning across the entire Israeli economy. Already in "
        "\u201cpreliminary operation,\u201d it will fuse data from multiple sources, "
        "classify threats using AI, and coordinate national defense."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The strategy describes the concept and states it is in preliminary operation. "
        "No performance metrics, detection rates, false positive rates, or operational "
        "data are provided. A German-Israeli joint pilot fusion center was announced "
        "for Q1 2026."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the Cyber Dome fails to achieve reliable national coverage, "
        "Israel\u2019s strategy loses its technological force multiplier. Defensive "
        "resources scatter across conventional tools without the centralized "
        "intelligence picture the strategy depends on. The NSOC and sectoral SOCs "
        "lose their primary data source, and the \u201cjoint defense\u201d concept "
        "has no integrating platform."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (55%): ")
    run.bold = True
    p.add_run(
        "We estimate a 55% probability that the Cyber Dome will demonstrate publicly "
        "documented effectiveness in preventing or mitigating a major cyber attack "
        "by end of 2028. Israel\u2019s high attack volume (26,000+ incidents in 2025, "
        "up 55%) creates many opportunities for documented success, and the INCD has "
        "strong institutional and budgetary incentives to publicize results. The "
        "German-Israeli cooperation agreement adds credibility pressure. However, "
        "no country has successfully deployed a nationwide AI cyber defense umbrella, "
        "operational security concerns may prevent specific attribution to the system, "
        "and the Iron Dome analogy (concept to operational effectiveness in ~5 years) "
        "suggests Cyber Dome may not reach full capability until 2027\u20132028."
    )

    # --- ASSUMPTION 2: Workforce ---
    doc.add_heading(
        "Assumption 2: The Cyber Workforce Can Scale to Meet Demand", level=2
    )

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "Israel will vigorously expand its cybersecurity workforce through education "
        "from elementary school through university, integration of underrepresented "
        "populations, breaking the \u201cjuniors\u2019 barrier,\u201d and employment "
        "of wounded veterans from the Iron Swords conflict."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Israel\u2019s cyber sector absorbs 36% of total hi-tech investment (2024). "
        "The military pipeline (Unit 8200 and equivalents) produces skilled graduates. "
        "But the strategy provides no workforce size data, growth targets, or analysis "
        "of the gap between supply and demand."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the workforce gap persists or widens, Israel\u2019s entire cyber "
        "posture degrades as threats grow faster than available defenders. The Cyber "
        "Dome needs skilled operators. The sectoral SOCs need staffing. The legislative "
        "framework creates compliance obligations that organizations cannot meet without "
        "qualified personnel. A workforce shortage is not just a constraint\u2014it is a "
        "ceiling on every other objective in the strategy."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (13%): ")
    run.bold = True
    p.add_run(
        "We estimate a 13% probability that Israel\u2019s cybersecurity workforce "
        "will grow by 30% from 2024 levels by end of 2028. Israel\u2019s overall "
        "tech workforce is currently shrinking for the first time in a decade (down "
        "1\u20132% annually), driven by war-related instability and salary pressures "
        "pushing firms to expand offshore. The cybersecurity pipeline produces "
        "approximately 2,800 graduates per year, but only ~40% enter the domestic "
        "private sector. The math is tight even under optimistic assumptions. "
        "Education initiatives targeting elementary students will not produce "
        "workers within the 2028 timeframe. The \u201cjuniors\u2019 barrier\u201d\u2014"
        "where 38% of hiring managers require 5+ years experience for entry-level "
        "roles\u2014is a persistent structural constraint that government programs "
        "have not overcome elsewhere."
    )

    # --- ASSUMPTION 3: Joint defense / legislation ---
    doc.add_heading(
        "Assumption 3: Joint Defense Will Overcome Regulatory Fragmentation",
        level=2,
    )

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "New dedicated cybersecurity legislation will regulate \u201cessential "
        "organizations,\u201d a National SOC will integrate sectoral centers, and "
        "the public-private \u201cwhole-of-society\u201d model will create coherent "
        "national defense. The strategy explicitly acknowledges that essential services "
        "organizations are \u201cnot currently adequately regulated.\u201d"
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The strategy references existing CERT operations, inter-agency coordination "
        "during Iron Swords, and the Privacy Security Law amendment (2025). A draft "
        "cybersecurity bill was published in January 2026. But the strategy does not "
        "address the INCD\u2019s two prior failed attempts to pass this legislation "
        "(2018 and 2021)."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that regulatory fragmentation persists\u2014legislation fails again, "
        "essential organizations remain outside mandatory frameworks\u2014the "
        "strategy\u2019s most actionable pillar collapses. The NSOC lacks authority "
        "to compel incident reporting from essential organizations. Sectoral SOCs "
        "operate without common standards. The \u201cwhole-of-society\u201d model "
        "remains aspirational rather than operational."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (25%): ")
    run.bold = True
    p.add_run(
        "We estimate a 25% probability that Israel will enact dedicated national "
        "cybersecurity legislation regulating essential organizations by end of 2027. "
        "A draft bill was published in January 2026\u2014the furthest this effort has "
        "ever progressed\u2014but elections scheduled for late 2026 will reset the "
        "legislative calendar. The INCD\u2019s two prior attempts died precisely "
        "because of electoral disruption. Even if a new government re-introduces the "
        "bill post-election, coalition formation delays, competing priorities (Haredi "
        "draft exemption, war-related legislation), and substantive opposition from "
        "privacy advocates compress the remaining legislative window to roughly "
        "9\u201315 months. Israel\u2019s coalition system, which has produced five "
        "elections in four years (2019\u20132022), makes sustained legislative "
        "commitment to any single bill uncertain."
    )

    # --- ASSUMPTION 4: International partnerships ---
    doc.add_heading(
        "Assumption 4: International Partnerships Will Remain Robust", level=2
    )

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "Israel will maintain and expand cybersecurity partnerships through Abraham "
        "Accords nations, OECD and INTERPOL engagement, tech company relationships, "
        "and intelligence sharing with friendly nations."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The Crystal Ball ransomware platform with UAE, Counter Ransomware Initiative "
        "membership, intelligence sharing with dozens of countries. Concrete existing "
        "partnerships with operational track records."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that partnerships erode significantly, Israel loses access to shared "
        "threat intelligence, coordinated vulnerability disclosure, and joint "
        "enforcement capabilities. However, the impact is mitigated by Israel\u2019s "
        "strong domestic capabilities and the bilateral (rather than multilateral) "
        "nature of most partnerships."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (95%): ")
    run.bold = True
    p.add_run(
        "We estimate a 95% probability that Israel will maintain or expand its "
        "international cybersecurity partnerships through end of 2028. This forecast "
        "is effectively already resolved: Germany signed a cyber defense cooperation "
        "agreement in January 2026, and India formalized a cybersecurity MoU in "
        "November 2025. Cybersecurity cooperation operates at the technical/agency "
        "level and is driven by pragmatic shared interest in defending against common "
        "threats, making it resilient to broader diplomatic friction. Israel\u2019s "
        "world-class cyber industry ($72.6 billion in exits) makes it a valued "
        "partner regardless of politics. No country has cancelled or suspended a "
        "cybersecurity cooperation agreement with Israel despite broader geopolitical "
        "tensions."
    )

    # --- ASSUMPTION 5: Iran threat ---
    doc.add_heading(
        "Assumption 5: The Iran-Centric Threat Model Is Adequate", level=2
    )

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "The strategy frames Iran and its proxies (Hezbollah, Hamas-affiliated groups) "
        "as the primary cyber adversary, with the October 7 / Iron Swords experience "
        "as the driving context. Other threats are mentioned but not emphasized."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Attack volumes doubled during Iron Swords, primarily from Iran-linked groups. "
        "12 billion NIS in annual cyber damage. But no comparative threat assessment "
        "of Russian, Chinese, or criminal capabilities targeting Israel."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that non-Iranian threats (Russia, China, criminal syndicates) become "
        "the primary cyber threat to Israel, the strategy\u2019s defensive architecture "
        "is calibrated to the wrong adversary. Russian APT operations are significantly "
        "more sophisticated than Iranian ones. If Russia-Israel relations deteriorate "
        "further, the threat profile shifts in ways the strategy doesn\u2019t address."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (87%): ")
    run.bold = True
    p.add_run(
        "We estimate an 87% probability that Iran and its proxies will remain the "
        "primary source (>50% of attributed attacks) of state-sponsored cyber attacks "
        "against Israel through end of 2028. Iran\u2019s dominance is structural: "
        "35+ active groups, 64% of all Iranian state-linked cyber activity directed "
        "at Israel, and operations have intensified 700% post-ceasefire. Russia-Iran "
        "cyber cooperation agreements are strengthening Iranian capabilities rather "
        "than creating a competing Russian threat. Pro-Russian hacktivists target "
        "Israel but primarily through low-sophistication DDoS. China\u2019s focus "
        "is IP theft rather than the disruptive operations that dominate attribution "
        "reports."
    )

    # --- ASSUMPTION 6: Timeline ---
    doc.add_heading(
        "Assumption 6: The 2028 Implementation Timeline Is Achievable", level=2
    )

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        'The strategy states that "a national implementation plan will be derived '
        "from this National Strategy Document, targeted to be completed by the next "
        'three years, namely, until 2028."'
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "None. The 2028 date appears to follow institutional planning cycles rather "
        "than an analysis of what is achievable. The strategy itself acknowledges "
        "that cyberspace is \u201cdynamic\u201d and will need frequent updates, and "
        "at the time of publication the detailed implementation plan did not yet exist."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the 2028 timeline is insufficient\u2014as the individual "
        "assumption forecasts suggest\u2014Israel operates under the old fragmented "
        "model during the period of highest threat escalation post-Iron Swords. "
        "This is the conjunction problem: meeting the deadline requires legislation "
        "(25%), workforce scaling (13%), and sustained political commitment "
        "simultaneously. Even the Cyber Dome (55%) is not a certainty. Planning "
        "should assume a 2030\u20132032 horizon for full implementation."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (25%): ")
    run.bold = True
    p.add_run(
        "We estimate a 25% probability that the majority (>60%) of the strategy\u2019s "
        "objectives will be completed by end of 2028. Israel will almost certainly "
        "make meaningful progress on several objectives (Cyber Dome, international "
        "partnerships, some institutional structures). But the legislative impasse, "
        "workforce constraints, and political instability\u2014elections are near-certain "
        "in 2026\u2014make the full timeline unrealistic. The 2017 national cyber "
        "strategy had \u201cmixed implementation results\u201d after 8 years, and "
        "international comparisons (US NIST framework: ~5 years to widespread adoption; "
        "EU NIS2: ~2 years proposal to adoption in a stable environment) suggest that "
        "three years is ambitious even without Israel\u2019s particular political "
        "constraints."
    )

    doc.add_page_break()

    # ========================================================
    # IMPLICATIONS FOR DECISION-MAKERS
    # ========================================================
    doc.add_heading("Implications for Decision-Makers", level=1)

    doc.add_heading("What to act on now", level=2)

    bullets_act = [
        (
            "Israel\u2019s threat assessment is reliable\u2014plan accordingly (87% Iran remains primary). ",
            "Organizations operating in or connected to the Israeli market should calibrate "
            "their threat models to Iranian cyber capabilities and tactics. Iran\u2019s "
            "operations are intensifying, not declining. Post-ceasefire cyber activity "
            "surged 700%. Expect continued escalation targeting critical infrastructure, "
            "supply chains, and influence operations."
        ),
        (
            "International partnerships are secure (95%)\u2014leverage them. ",
            "Israel\u2019s cybersecurity cooperation framework is expanding despite "
            "geopolitical headwinds. Organizations can rely on continued intelligence "
            "sharing, coordinated vulnerability disclosure, and frameworks like the "
            "Counter Ransomware Initiative. The Abraham Accords cyber cooperation "
            "offers new regional opportunities."
        ),
        (
            "Plan for a workforce-constrained environment (87% the gap persists). ",
            "Israel\u2019s tech workforce is shrinking for the first time in a decade. "
            "Organizations should assume persistent difficulty hiring qualified "
            "cybersecurity personnel in Israel and invest in automation, managed "
            "security services, and retention strategies rather than expecting the "
            "pipeline to expand."
        ),
    ]

    for bold_text, normal_text in bullets_act:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_heading("What to watch", level=2)

    bullets_watch = [
        (
            "Cybersecurity legislation (currently 25% by end of 2027). ",
            "The draft bill published January 2026 is the most concrete step yet. "
            "If it passes before or shortly after the expected 2026 elections, it "
            "would transform the regulatory landscape for essential organizations. "
            "Monitor Knesset committee progress and post-election government "
            "priorities."
        ),
        (
            "Cyber Dome operational milestones (currently 55%). ",
            "Track INCD public disclosures about specific incidents where the Cyber "
            "Dome system contributed to defense. The German-Israeli fusion center "
            "(expected 24/7 operations Q1 2026) is an early indicator. If the INCD "
            "begins citing Cyber Dome in specific incident reports, the system is "
            "maturing."
        ),
        (
            "Post-election government formation (expected late 2026). ",
            "The new government\u2019s priorities will determine whether the strategy "
            "maintains momentum or stalls. Watch for: cybersecurity legislation in "
            "the coalition agreement, INCD budget allocations, and whether the new "
            "INCD chief (if appointed) maintains the current strategic direction."
        ),
    ]

    for bold_text, normal_text in bullets_watch:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_page_break()

    # ========================================================
    # METHODOLOGY NOTE
    # ========================================================
    doc.add_heading("Methodology", level=1)

    doc.add_paragraph(
        "Each assumption was converted into a concrete, time-bound forecasting question "
        "and evaluated using structured forecasting methods:"
    )

    for step in [
        "Base rate construction from historical reference classes",
        "Evidence gathering from domain research and recent publications",
        "Bayesian updating based on specific factors for and against",
        "Decomposition analysis (breaking the question into independent paths to resolution)",
        "Calibration checks and stress testing against sensitivity ranges",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph(
        "Forecasts represent calibrated probabilities\u2014our best estimate of likelihood "
        "given available evidence. A 60% forecast means we expect the event to occur in "
        "roughly 6 out of 10 similar situations. These are not expressions of confidence "
        "in our analysis; they are estimates of real-world probability."
    )

    doc.add_paragraph(
        'Sensitivity framing uses presuppositional ("Given that X") rather than '
        'hypothetical ("If X were true") language. Research on conditional reasoning '
        "shows that presuppositional framing produces more accurate counterfactual "
        "analysis by treating the alternative scenario as a concrete state of the world "
        "rather than an abstract possibility."
    )

    # ========================================================
    # SAVE
    # ========================================================
    output_path = "/Users/elsehow/Downloads/audit/javier/FRI_Forecast_Audit_Israel_Cybersecurity_Strategy_2025.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
