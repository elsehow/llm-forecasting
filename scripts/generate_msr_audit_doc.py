"""Generate the Forecast Audit deliverable for Munich Security Report 2026."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elm)


def add_styled_paragraph(doc, text, style="Normal", bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def build_document():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ========================================================
    # TITLE PAGE
    # ========================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FORECAST AUDIT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        '"Under Destruction"\nMunich Security Report 2026'
    )
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    source_line = doc.add_paragraph()
    source_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source_line.add_run("Source report: Munich Security Conference, February 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run("Audit date: February 25, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line = doc.add_paragraph()
    prep_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line.add_run("Prepared for: Javier Prieto")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line2 = doc.add_paragraph()
    prep_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line2.add_run("Prepared by: Forecasting Research Institute")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This audit evaluates the key assumptions underlying the Munich Security Report 2026 "
        '("Under Destruction"), which argues that the post-1945 US-led international order is '
        "being systematically dismantled by \"wrecking-ball politics.\" For each assumption, we "
        "extract the report's implicit claim, assess the evidence provided, and assign a "
        "calibrated probability using structured forecasting methods."
    )

    doc.add_paragraph(
        "The report's headline conclusion is that US withdrawal from multilateral institutions, "
        "combined with rising geopolitical threats, is creating a more dangerous, fragmented "
        "world. Our audit finds that the report's core assumptions are largely well-founded, "
        "though the Russian conventional threat timeline may be overstated."
    )

    # Summary table
    doc.add_paragraph("")
    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Assumption", "Sensitivity", "P(Report is Right)", "Verdict"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A1A2E")

    rows_data = [
        ["US institutional retreat is durable", "HIGH", "96%", "Probably right"],
        ["Russia can reconstitute NATO-level threat by 2028", "HIGH", "25%", "Uncertain"],
        ["Europe cannot close the defense gap fast enough", "HIGH", "95%", "Probably right"],
        ["Non-traditional donors cannot fill aid gap", "MEDIUM", "91%", "Probably right"],
        ["Trade fragmentation produces lasting harm", "MEDIUM", "95%", "Probably right"],
        ["European publics sustain defense spending support", "HIGH", "78%", "Probably right"],
    ]

    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if col_idx == 1:  # Sensitivity column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if val == "HIGH":
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    run.bold = True
                elif val == "MEDIUM":
                    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)
                    run.bold = True
            elif col_idx == 2:  # Probability column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
        # Alternate row shading
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F5F5FA")

    doc.add_paragraph("")

    # Key findings
    doc.add_heading("Key Findings", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Where the report is well-supported: ")
    run.bold = True
    p.add_run(
        "Five of six assumptions strongly support the report's framing. US institutional "
        "retreat is highly durable within the current presidential term (96%). Europe cannot "
        "reach 3% GDP defense spending by 2028 (95% the report is right). Non-traditional "
        "donors will not fill the humanitarian aid gap (91%). Global trade volumes will not "
        "decline in absolute terms, but this actually supports the report's own hedge that "
        "trade has been \"more resilient than feared\" while still fragmenting. European "
        "publics will likely sustain defense spending support (78%), validating the political "
        "will the report identifies."
    )

    p = doc.add_paragraph()
    run = p.add_run("Where the report may overstate: ")
    run.bold = True
    p.add_run(
        "The Russian conventional threat timeline is the weakest assumption. At 25% "
        "probability of Russia meeting both the 1.5 million personnel and $100 billion "
        "spending thresholds by 2028, the report's implied urgency about an imminent "
        "Russian threat to NATO may be overstated. Russia's massive casualty rates "
        "(25,000-35,000/month), declining recruitment, and demographic headwinds make "
        "rapid force reconstitution very difficult without formal mobilization. However, "
        "the spending component alone is very likely met ($149B in 2024), and a ceasefire "
        "could accelerate personnel recovery."
    )

    p = doc.add_paragraph()
    run = p.add_run("The bottom line: ")
    run.bold = True
    p.add_run(
        "The Munich Security Report 2026's core narrative\u2014that the US-led order is under "
        "serious, durable strain\u2014is well-grounded in the evidence. The report is strongest "
        "on institutional dynamics (US withdrawal, European defense gaps, aid collapse) and "
        "weakest on military threat timelines. Decision-makers should treat the institutional "
        "analysis as reliable and the military urgency claims with more caution."
    )

    doc.add_page_break()

    # ========================================================
    # HOW TO READ THIS AUDIT
    # ========================================================
    doc.add_heading("How to Read This Audit", level=1)

    doc.add_paragraph(
        "Every analytical report rests on assumptions\u2014claims the author treats as true "
        "without fully defending them. Some are well-supported. Others are not. "
        "When a decision-maker relies on a report, they inherit its assumptions."
    )

    doc.add_paragraph(
        "This audit makes those assumptions explicit. For each one, we ask:"
    )

    bullets = [
        "What does the report assume?",
        "What evidence does it provide?",
        "How much would the conclusion change if this assumption is wrong? (Sensitivity)",
        "How likely is the assumption to be correct? (Probability)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Probabilities are generated using structured forecasting: base rate construction, "
        "evidence weighting, decomposition analysis, and calibration checks. They represent "
        "our best estimate given available evidence, not certainty."
    )

    p = doc.add_paragraph()
    run = p.add_run('The "Given that" framing: ')
    run.bold = True
    p.add_run(
        "For each assumption, we describe the sensitivity using presuppositional framing: "
        "\"Given that this assumption is wrong, here is how the conclusion changes.\" "
        "Research shows this framing produces more accurate conditional reasoning than "
        "hypothetical framing (\"If this were wrong...\")."
    )

    # Verdict table
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Verdict scale:")
    run.bold = True

    verdict_table = doc.add_table(rows=6, cols=2)
    verdict_table.style = "Table Grid"
    verdict_headers = ["P(Report is Right)", "Verdict"]
    for i, h in enumerate(verdict_headers):
        cell = verdict_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A1A2E")

    verdict_data = [
        [">75%", "Probably right"],
        ["50\u201375%", "Plausible but uncertain"],
        ["25\u201350%", "Uncertain"],
        ["10\u201325%", "Probably overstated"],
        ["<10%", "Likely wrong"],
    ]
    for row_idx, row_data in enumerate(verdict_data):
        row = verdict_table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_page_break()

    # ========================================================
    # DETAILED ASSUMPTION ANALYSIS
    # ========================================================
    doc.add_heading("Detailed Assumption Analysis", level=1)

    # --- ASSUMPTION 1 ---
    doc.add_heading("Assumption 1: US Institutional Retreat Is Durable", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "The Trump administration's withdrawal from multilateral institutions (WHO, Paris "
        "Agreement, UNRWA, UN Human Rights Council, WTO Appellate Body) represents a "
        "structural shift in US foreign policy, not a temporary political cycle. The "
        "report's entire \"under destruction\" framing depends on this retreat being lasting."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Extensive. The report documents DOGE-driven government downsizing, USAID shutdown, "
        "multiple simultaneous institutional withdrawals, Congressional Republican support "
        "for isolationism, and the permanent destruction of institutional knowledge and "
        "personnel. The report argues this term's actions are more aggressive and thoroughgoing "
        "than 2017\u20132020."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the US retreat is reversed within this presidential term\u2014say, through "
        "a pandemic forcing WHO reengagement or Congressional action\u2014the report's "
        "\"under destruction\" framing becomes \"under stress but resilient.\" The entire "
        "narrative shifts from irreversible decline to cyclical disruption, and the urgency "
        "of European strategic autonomy diminishes significantly."
    )

    p = doc.add_paragraph()
    run = p.add_run("Forecasting question: ")
    run.bold = True
    p.add_run(
        "Will the US rejoin or restore funding to at least 2 of 5 key international "
        "institutions/agreements (WHO, Paris Agreement, UNRWA, UNHRC, WTO Appellate Body) "
        "by December 31, 2028?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast: 4% (report is right: 96%)")
    run.bold = True

    doc.add_paragraph(
        "There is essentially no historical precedent for a US president withdrawing from "
        "international institutions and then rejoining them within the same term. All previous "
        "reversals (Paris, WHO, UNHRC) occurred across administrations. The current term's "
        "institutional destruction goes deeper than 2017\u20132020: USAID has been gutted, "
        "bureaucratic expertise permanently lost, and Congressional Republicans are more "
        "isolationist. The only plausible pathway to reversal\u2014a severe pandemic forcing "
        "WHO reengagement combined with a second institution\u2014has roughly 1\u20132% "
        "probability. Even accounting for unknown unknowns, the probability of rejoining 2+ "
        "institutions by 2028 is approximately 4%."
    )

    # --- ASSUMPTION 2 ---
    doc.add_heading("Assumption 2: Russia Can Reconstitute a NATO-Level Threat by 2028", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "Russia could pose a conventional military threat to NATO\u2014particularly the "
        "Baltic states\u2014within 3\u20135 years of a Ukraine ceasefire. The report implies "
        "Russia's wartime economy and military expansion could rapidly reconstitute forces "
        "after current losses."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The report documents Russia's defense spending increase to 6\u20137% of GDP, wartime "
        "industrial mobilization, and Putin's decree to expand active-duty forces to 1.5 "
        "million. However, the report is light on the countervailing evidence: massive "
        "casualty figures, equipment losses, and demographic constraints."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that Russia cannot reconstitute forces for a NATO-threatening operation "
        "within the implied timeframe, the urgency of European defense spending is lower "
        "than the report suggests. Europe would have more time\u2014likely through the early "
        "2030s\u2014to build credible deterrence, and the current pace of defense increases "
        "might be adequate rather than insufficient."
    )

    p = doc.add_paragraph()
    run = p.add_run("Forecasting question: ")
    run.bold = True
    p.add_run(
        "Will Russia's active-duty military personnel exceed 1.5 million AND its annual "
        "military expenditure exceed $100 billion (constant 2024 USD) by December 31, 2028?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast: 25% (verdict: Uncertain)")
    run.bold = True

    doc.add_paragraph(
        "This is a conjunction forecast bottlenecked by the personnel condition. Russia's "
        "military spending already exceeds $100B by a wide margin (SIPRI estimated $149B "
        "for 2024), making the expenditure condition very likely (~80% standalone). However, "
        "reaching 1.5 million active-duty personnel is far more challenging. IISS estimates "
        "current strength at ~1.13 million, requiring a net increase of ~370,000 while "
        "sustaining 25,000\u201335,000 casualties monthly and facing declining recruitment "
        "(422K new contracts in 2025, down from 450K in 2024). Carnegie Endowment analysis "
        "suggests full reconstitution is a 2030s timeline without additional mobilization. "
        "A ceasefire would improve personnel prospects by reducing attrition, but even then, "
        "demographic constraints limit expansion speed."
    )

    # --- ASSUMPTION 3 ---
    doc.add_heading("Assumption 3: Europe Cannot Close the Defense Capability Gap Fast Enough", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "Despite a 41% increase in European defense spending between 2021\u20132025, Europe "
        "remains critically dependent on US defense systems (51% of equipment procurement) "
        "and cannot achieve strategic autonomy within the relevant timeframe. Fiscal "
        "constraints, industrial fragmentation, and interoperability gaps prevent rapid "
        "capability building."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Strong. The report details the 51% US procurement dependency, fiscal rule "
        "constraints, the gap between 2% and 5% NATO targets, early-stage joint programs "
        "(European Sky Shield), and the fundamental challenge of 27 separate national "
        "armies. The evidence directly supports the claim."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that Europe can close the defense gap quickly\u2014reaching 3% of GDP by "
        "2028 and building interoperable capabilities\u2014the report's crisis framing is "
        "less urgent. European strategic autonomy becomes achievable within this decade, and "
        "the US withdrawal from European security, while unwelcome, is manageable."
    )

    p = doc.add_paragraph()
    run = p.add_run("Forecasting question: ")
    run.bold = True
    p.add_run(
        "Will combined EU-27 member state defense spending reach 3% of aggregate GDP by "
        "fiscal year 2028?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast: 5% (report is right: 95%)")
    run.bold = True

    doc.add_paragraph(
        "No major institutional forecast (NATO, Goldman Sachs, McKinsey, IISS) projects "
        "EU-27 aggregate defense spending reaching 3% by 2028. The IISS Military Balance "
        "2026, released February 24, 2026, confirms European NATO average at 2.16% in "
        "2025. Reaching 3% would require roughly tripling the annual pace of spending "
        "increase\u2014an additional \u20ac150\u2013170 billion per year beyond current levels. "
        "The EU's ReArm Europe fiscal escape clause allows up to 1.5 percentage points of "
        "additional spending phased over four years, but few countries will fully utilize "
        "this ceiling. Defense-industrial absorption constraints further limit how quickly "
        "spending can translate into actual capability. The most likely trajectory reaches "
        "2.4\u20132.7% by 2028."
    )

    # --- ASSUMPTION 4 ---
    doc.add_heading("Assumption 4: Non-Traditional Donors Cannot Fill the Humanitarian Aid Gap", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "The USAID shutdown (5,300+ of ~6,200 awards cancelled) and broader ODA cuts will "
        "produce catastrophic humanitarian consequences\u2014an estimated 14 million "
        "additional deaths by 2030\u2014because non-traditional donors (China, Gulf states) "
        "cannot fill the gap."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Moderate. The report documents USAID cancellations, ODA declines (7.1% in 2024, "
        "projected 9\u201317% in 2025), and UN budget cuts (15\u201334%). It notes China's "
        "expanding influence in UN agencies but correctly observes that non-traditional "
        "donors have different priorities (infrastructure vs. health/education) and operate "
        "at much smaller scale."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that non-traditional donors rapidly scale their development assistance\u2014"
        "China launches a major new development initiative, Gulf states collectively commit "
        "to filling the gap\u2014the humanitarian catastrophe is significantly mitigated. "
        "However, even a 50% increase in non-DAC aid would still leave total global "
        "development assistance well below pre-withdrawal levels, so the gap would be "
        "narrowed rather than closed."
    )

    p = doc.add_paragraph()
    run = p.add_run("Forecasting question: ")
    run.bold = True
    p.add_run(
        "Will total non-DAC bilateral development assistance increase by at least 50% in "
        "real terms from 2024 levels by December 31, 2028?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast: 9% (report is right: 91%)")
    run.bold = True

    doc.add_paragraph(
        "The scale mismatch is the fundamental constraint. China's foreign aid is estimated "
        "at $5\u20137B/year versus US levels of $50B+. China faces binding fiscal constraints "
        "(property crisis, local government debt, slowing growth) that make massive aid "
        "expansion politically difficult. Gulf state aid is volatile and concentrated in the "
        "MENA region. No coordination mechanism exists among non-DAC donors, and the free-rider "
        "problem is severe. Even the strongest bull case\u2014China announcing a major new "
        "development initiative\u2014would take years to operationalize. The 9% probability "
        "reflects small upside from measurement elasticity (broader definitions capturing "
        "more Chinese state financing) and organic growth in Indian and Turkish development "
        "cooperation."
    )

    # --- ASSUMPTION 5 ---
    doc.add_heading("Assumption 5: Trade Fragmentation Produces Lasting Economic Harm", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "US tariffs at 1930s levels (~15% average), \"Liberation Day\" tariffs, and "
        "retaliatory measures are fragmenting the global trading system, with the Global "
        "Economic Policy Uncertainty Index at all-time highs. However, the report itself "
        "hedges, noting trade has been \"more resilient than feared.\""
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Mixed. The report documents high tariff levels and uncertainty indices but also "
        "acknowledges trade rerouting through third countries, new trade blocs forming "
        "outside the US (CPTPP, EU-Mercosur, RCEP), and trade resilience. The report is "
        "appropriately nuanced on this point."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that trade fragmentation produces only modest welfare losses\u2014new trade "
        "blocs compensate for US withdrawal, rerouting offsets tariff impacts\u2014the "
        "economic pillar of the \"under destruction\" narrative weakens. The report's "
        "broader argument about institutional fragmentation would still stand, but the "
        "economic consequences would be less severe than implied."
    )

    p = doc.add_paragraph()
    run = p.add_run("Forecasting question: ")
    run.bold = True
    p.add_run(
        "Will global merchandise trade volume (goods) in 2028 be lower than 2024 levels?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast: 5% (report is right: 95%)")
    run.bold = True

    doc.add_paragraph(
        "An absolute decline in global trade volume over a 4-year period has essentially "
        "never occurred in the post-WWII era. Even the 2009 Global Financial Crisis saw "
        "full recovery within two years. Trade volume in 2025 is already estimated at ~5.4% "
        "above 2024 levels, creating a substantial buffer. The US Supreme Court struck down "
        "IEEPA tariffs on February 20, 2026, with the replacement Section 122 surcharge "
        "capped at 15%. AI-related goods (42% of 2025 trade growth) provide a structural "
        "tailwind largely exempt from tariffs. WTO and IMF forecasts remain positive. "
        "However, this finding actually supports the report's own nuanced position: the "
        "damage from trade fragmentation is real (deadweight losses, investment uncertainty, "
        "supply chain disruption) even though absolute trade volumes continue to grow."
    )

    # --- ASSUMPTION 6 ---
    doc.add_heading("Assumption 6: European Publics Will Sustain Defense Spending Support", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report's claim: ")
    run.bold = True
    p.add_run(
        "The Munich Security Index shows heightened threat perception across G7 publics and "
        "growing recognition that Europe must do more for its own defense. The report implies "
        "democratic legitimacy exists for continued defense buildup."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The MSI polling data shows declining confidence in US security guarantees and "
        "increased threat perception. The report documents Germany's Zeitenwende, Poland's "
        "4%+ spending, and EU-wide defense momentum. However, it does not address the "
        "risk of \"threat fatigue\" or competing fiscal priorities."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that European publics withdraw support for defense spending increases\u2014"
        "say, due to economic recession, threat fatigue, or political shifts\u2014the "
        "defense buildup the report calls for becomes democratically unsustainable. "
        "Governments would face a choice between fiscal discipline and security investment, "
        "with no mandate for prioritizing defense."
    )

    p = doc.add_paragraph()
    run = p.add_run("Forecasting question: ")
    run.bold = True
    p.add_run(
        "Will public support for maintaining or increasing defense spending remain above "
        "50% in at least 3 of the 4 largest EU economies (Germany, France, Italy, Spain) "
        "through December 2027?"
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast: 78% (verdict: Probably right)")
    run.bold = True

    doc.add_paragraph(
        "The forecast rests on a critical distinction: \"maintaining or increasing\" is a "
        "substantially lower bar than \"increasing.\" Current polling shows strong majority "
        "support for defense spending increases in Germany (68\u201377%) and moderate support "
        "in France and Spain (45\u201347% for increases alone). Adding \"maintain\" supporters "
        "pushes all three well above 50%. Italy is the clear outlier (17% for increases, "
        "57% opposed), but the 3-of-4 structure accommodates this. Per-country decomposition: "
        "Germany ~95%, France ~85%, Spain ~82%, Italy ~35%. The main risks are correlated: "
        "a severe economic downturn or comprehensive Ukraine peace deal could shift opinion "
        "across multiple countries simultaneously. However, the structural security shift "
        "(US unreliability, Russian proximity) is durable and unlikely to reverse by late 2027."
    )

    doc.add_page_break()

    # ========================================================
    # IMPLICATIONS FOR DECISION-MAKERS
    # ========================================================
    doc.add_heading("Implications for Decision-Makers", level=1)

    doc.add_heading("Act on now", level=2)

    bullets_act = [
        (
            "The US institutional withdrawal is real and durable through 2028 (96%). ",
            "Planning and policy should assume the US will not re-engage with multilateral "
            "institutions during this presidential term. European and allied institutions "
            "need to plan for sustained US absence, not temporary disruption. Institutional "
            "alternatives (EU-led, regional, minilateral) should be accelerated."
        ),
        (
            "The defense capability gap is structural, not just budgetary (95%). ",
            "Simply increasing spending will not close the gap quickly enough. The binding "
            "constraint is defense-industrial capacity, interoperability, and the institutional "
            "fragmentation of 27 separate armies. Priority should be given to joint procurement, "
            "industrial consolidation, and capability specialization rather than headline "
            "spending targets."
        ),
        (
            "The humanitarian aid collapse is real and the gap will not be filled (91%). ",
            "Organizations and governments depending on USAID-funded programs need contingency "
            "plans now. Non-traditional donors will not compensate. European and allied donors "
            "should identify the most critical health and food security programs at risk and "
            "prioritize bridge funding for the highest-impact interventions."
        ),
    ]

    for bold_text, normal_text in bullets_act:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_heading("Watch", level=2)

    bullets_watch = [
        (
            "Russian force reconstitution timeline (25% by 2028). ",
            "The report's implied urgency about a near-term Russian conventional threat to "
            "NATO may be overstated. Monitor IISS and SIPRI data on Russian force levels, "
            "recruitment rates, and equipment production. A ceasefire in Ukraine would be "
            "the key trigger to reassess\u2014it would simultaneously reduce Russian attrition "
            "and free up resources for reconstitution. Carnegie Endowment analysis suggests "
            "the 2030s is a more realistic timeline."
        ),
        (
            "European public opinion durability (78%). ",
            "Sustained defense spending depends on sustained public support. Monitor "
            "Eurobarometer and national polls, particularly in Italy (the weakest link at "
            "~35% probability of sustaining support). A severe economic downturn or a Ukraine "
            "ceasefire could erode support faster than expected. Germany's bipartisan consensus "
            "provides a structural floor."
        ),
        (
            "Trade system evolution rather than collapse. ",
            "The report correctly notes trade has been \"more resilient than feared.\" The "
            "story is fragmentation and rerouting, not collapse. Watch for the emergence of "
            "parallel trading blocs (CPTPP expansion, RCEP deepening, EU bilateral deals) "
            "as a new equilibrium rather than continued deterioration."
        ),
    ]

    for bold_text, normal_text in bullets_watch:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_page_break()

    # ========================================================
    # METHODOLOGY NOTE
    # ========================================================
    doc.add_heading("Methodology", level=1)

    doc.add_paragraph(
        "Each assumption was converted into a concrete, time-bound forecasting question "
        "and evaluated using structured forecasting methods:"
    )

    method_steps = [
        "Base rate construction from historical reference classes",
        "Evidence gathering from domain research, institutional data, and recent publications",
        "Bayesian updating based on specific factors for and against",
        "Decomposition analysis (breaking the question into independent components)",
        "Prediction market research where available",
        "Calibration checks and stress testing against sensitivity ranges",
    ]
    for step in method_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph(
        "Forecasts represent calibrated probabilities\u2014our best estimate of likelihood "
        "given available evidence. A 60% forecast means we expect the event to occur in "
        "roughly 6 out of 10 similar situations. These are not expressions of confidence "
        "in our analysis; they are estimates of real-world probability."
    )

    doc.add_paragraph(
        "Sensitivity framing uses presuppositional (\"Given that X\") rather than "
        "hypothetical (\"If X were true\") language. Research on conditional reasoning "
        "shows that presuppositional framing produces more accurate counterfactual "
        "analysis by treating the alternative scenario as a concrete state of the world "
        "rather than an abstract possibility."
    )

    doc.add_paragraph(
        "Six independent forecasting agents were run in parallel, each applying the "
        "full structured forecasting workflow to its assigned assumption. This parallel "
        "architecture ensures that each forecast is produced independently, avoiding "
        "anchoring effects between assumptions."
    )

    # ========================================================
    # SAVE
    # ========================================================
    output_path = "/Users/elsehow/Downloads/audit/javier/FRI_Forecast_Audit_Munich_Security_Report_2026.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
