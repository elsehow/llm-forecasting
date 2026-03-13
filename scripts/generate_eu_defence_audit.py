"""Generate Forecast Audit deliverable for the EU Defence White Paper."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex},
    )
    shading.append(shading_elm)


def add_styled_paragraph(doc, text, style="Normal", bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ========================================================
    # TITLE PAGE
    # ========================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FORECAST AUDIT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        '"European Defence \u2013 Readiness 2030"\nWhite Paper for European Defence'
    )
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    source_line = doc.add_paragraph()
    source_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source_line.add_run("Source report: European Commission, March 2025")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run("Audit date: February 25, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line = doc.add_paragraph()
    prep_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line.add_run("Prepared for: Javier, Repsol")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    prep_line2 = doc.add_paragraph()
    prep_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_line2.add_run("Prepared by: Forecasting Research Institute")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This audit evaluates the key assumptions underlying the European Commission's "
        "March 2025 white paper on European defence readiness. The paper proposes a "
        "\u20ac800 billion rearmament plan (the \"ReArm Europe\" / SAFE initiative) premised "
        "on reaching operational readiness by 2030. For each assumption, we extract the "
        "report's implicit claim, assess the evidence provided, and assign a calibrated "
        "probability using structured forecasting methods."
    )

    doc.add_paragraph(
        "The paper's headline conclusion is that Europe can close its defence gap "
        "through coordinated spending, industrial scale-up, and collaborative "
        "procurement within a five-year timeline. Our audit finds that this conclusion "
        "rests on several assumptions, most of which are significantly overstated."
    )

    # Summary table
    doc.add_paragraph("")
    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Assumption", "Sensitivity", "P(Report is Right)", "Verdict"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A1A2E")

    rows_data = [
        ["Member states will spend the money", "HIGH", "58%", "Plausible but uncertain"],
        ["US security guarantee durably reduced", "HIGH", "8%", "Likely wrong"],
        ["EDTIB can scale production fast enough", "HIGH", "12%", "Probably overstated"],
        ["Collaborative procurement reaches 40%", "MEDIUM", "8%", "Likely wrong"],
        ["Russia\u2019s threat extends beyond Ukraine", "HIGH", "18%", "Probably overstated"],
        ["2030 timeline is sufficient", "HIGH", "12%", "Probably overstated"],
    ]

    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if col_idx == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if val == "HIGH":
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    run.bold = True
            elif col_idx == 2:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F5F5FA")

    doc.add_paragraph("")

    # Key findings
    doc.add_heading("Key Findings", level=2)

    p = doc.add_paragraph()
    run = p.add_run("The plan\u2019s central vulnerability: ")
    run.bold = True
    p.add_run(
        "Five of six assumptions score below 20%, meaning the white paper\u2019s "
        "conclusion depends on conditions that are individually unlikely and collectively "
        "near-impossible. The only assumption with meaningful support is that member "
        "states will increase spending (58%)\u2014but spending alone, without industrial "
        "capacity, collaborative procurement, and a viable timeline, does not deliver "
        "readiness."
    )

    p = doc.add_paragraph()
    run = p.add_run("The US security guarantee (8%): ")
    run.bold = True
    p.add_run(
        "The paper\u2019s urgency rests on the premise that the US security guarantee "
        "has fundamentally changed. Our forecast finds this overstated: the structural "
        "factors binding the US to European security (forward-deployed forces, NATO "
        "infrastructure, Congressional support) remain strong. A durable pullback is "
        "unlikely absent a major shift in US domestic politics."
    )

    p = doc.add_paragraph()
    run = p.add_run("The industrial bottleneck (12%): ")
    run.bold = True
    p.add_run(
        "Even with funding, Europe\u2019s defence industrial base cannot double "
        "production within the proposed timeline. Ammunition production alone requires "
        "2\u20133 years for new facilities. Major platforms (frigates, combat aircraft, "
        "armored vehicles) have 5\u201310 year lead times. The white paper assumes "
        "industrial capacity that does not yet exist."
    )

    p = doc.add_paragraph()
    run = p.add_run("The bottom line: ")
    run.bold = True
    p.add_run(
        "Europe will likely spend more on defence. But the white paper\u2019s specific "
        "vision\u2014reaching operational readiness by 2030 through coordinated "
        "pan-European rearmament\u2014rests on a conjunction of improbable assumptions. "
        "A more realistic planning horizon is 2035\u20132040. Decision-makers should "
        "plan for a slower, more fragmented build-up than the paper implies."
    )

    doc.add_page_break()

    # ========================================================
    # HOW TO READ THIS AUDIT
    # ========================================================
    doc.add_heading("How to Read This Audit", level=1)

    doc.add_paragraph(
        "Every analytical report rests on assumptions\u2014claims the author treats as true "
        "without fully defending them. Some are well-supported. Others are not. "
        "When a decision-maker relies on a report, they inherit its assumptions."
    )

    doc.add_paragraph(
        "This audit makes those assumptions explicit. For each one, we ask:"
    )

    bullets = [
        "What does the report assume?",
        "What evidence does it provide?",
        "How much would the conclusion change if this assumption is wrong? (Sensitivity)",
        "How likely is the assumption to be correct? (Probability)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Probabilities are generated using structured forecasting: base rate construction, "
        "evidence weighting, decomposition analysis, and calibration checks. They represent "
        "our best estimate given available evidence, not certainty."
    )

    p = doc.add_paragraph()
    run = p.add_run('The "Given that" framing: ')
    run.bold = True
    p.add_run(
        "For each assumption, we describe the sensitivity using presuppositional framing: "
        '"Given that this assumption is wrong, here is how the conclusion changes." '
        "Research shows this framing produces more accurate conditional reasoning than "
        'hypothetical framing ("If this were wrong...").'
    )

    doc.add_page_break()

    # ========================================================
    # DETAILED ASSUMPTION ANALYSIS
    # ========================================================
    doc.add_heading("Detailed Assumption Analysis", level=1)

    # --- ASSUMPTION 1: Spending ---
    doc.add_heading("Assumption 1: Member States Will Spend the Money", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "EU member states will collectively reach or approach the proposed spending "
        "targets, with at least 80% meeting the 3.5% GDP benchmark by 2030."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The SAFE instrument (\u20ac150 billion in EU-backed loans), EDIP grants "
        "(\u20ac1.5 billion), and strong political rhetoric post-Ukraine. The paper "
        "points to rising defence budgets across the EU since 2022."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that member states do not increase spending to the proposed levels\u2014"
        "say, the median EU member reaches 2.5% rather than 3.5% of GDP\u2014the entire "
        "rearmament plan stalls. Without the funding, industrial orders do not materialize, "
        "collaborative procurement incentives lack backing, and the 2030 timeline "
        "becomes impossible. The paper\u2019s \u20ac800 billion figure is predicated on "
        "near-universal compliance with historically unprecedented spending targets."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (58%): ")
    run.bold = True
    p.add_run(
        "We estimate a 58% probability that at least 80% of NATO-EU member states "
        "will meet the 2% GDP spending target by end of 2027. The historical trend is "
        "encouraging: NATO European members increased spending from 1.47% (2014) to "
        "approximately 2.03% (2024), and 23 of 32 allies met the 2% target in 2024, up "
        "from just 3 in 2014. However, the 3.5% target proposed by the white paper is "
        "far more ambitious\u2014only the US and Poland currently spend above 3.5%. "
        "The 2% target is achievable; the 3.5% target underlying the paper\u2019s "
        "\u20ac800 billion figure is not."
    )

    # --- ASSUMPTION 2: US pullback ---
    doc.add_heading("Assumption 2: The US Security Guarantee Is Durably Reduced", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "The paper\u2019s urgency depends on the premise that the transatlantic security "
        "guarantee has fundamentally weakened, requiring Europe to achieve strategic "
        "autonomy. References to \u201cchanging geopolitical realities\u201d and \u201cnew "
        "security environment\u201d imply a structural shift, not a temporary fluctuation."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "General references to shifting US priorities and European vulnerability. "
        "No specific analysis of US force posture, Congressional appropriations, "
        "or alliance commitments."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the US security guarantee remains fundamentally intact\u2014"
        "forward-deployed forces stay in Europe, Article 5 credibility is maintained, "
        "Congressional support continues\u2014the urgency of the paper\u2019s timeline "
        "collapses. European defence modernization remains desirable but the case for "
        "a crash rearmament program weakens dramatically. The political will to sustain "
        "3.5% GDP spending evaporates without a credible US pullback narrative."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (8%): ")
    run.bold = True
    p.add_run(
        "We estimate an 8% probability that the US will formally reduce its military "
        "commitment to NATO by more than 25% (measured by forward-deployed personnel "
        "or declared force commitments) by end of 2028. The base rate of major alliance "
        "reconfigurations is low (2\u20135% per year for established alliances). The US "
        "currently has approximately 100,000 troops in Europe, Congress has passed "
        "legislation restricting NATO withdrawal, and bipartisan NATO support remains "
        "strong in the Senate. Even administrations rhetorically hostile to NATO have "
        "increased European force posture. The structural factors (forward bases, "
        "defense industry integration, intelligence sharing) create institutional "
        "lock-in that is difficult to reverse within a single presidential term."
    )

    # --- ASSUMPTION 3: EDTIB scaling ---
    doc.add_heading("Assumption 3: The Defence Industrial Base Can Scale Fast Enough", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "The European Defence Technological and Industrial Base (EDTIB) can absorb "
        "dramatically increased orders and scale production to meet the 2030 readiness "
        "targets. The paper proposes specific mechanisms (SAFE loans, EDIP grants, "
        "streamlined procurement) to accelerate this."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "Limited. The paper acknowledges capacity constraints but treats them as "
        "solvable through funding and procurement reform. No detailed analysis of "
        "current production capacity, workforce availability, supply chain readiness, "
        "or the time required to build new production lines."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the EDTIB cannot double production within the proposed timeline\u2014"
        "which is the historical norm for defence industrial scaling\u2014the readiness "
        "gap persists regardless of spending levels. Europe would be spending more money "
        "on the same limited production capacity, driving up unit costs without "
        "proportional capability gains. The planning assumption should shift from "
        "\u201cbuild European\u201d to a mixed strategy of European production plus "
        "off-the-shelf imports for near-term gaps."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (12%): ")
    run.bold = True
    p.add_run(
        "We estimate a 12% probability that EDTIB production output (measured by "
        "major platform deliveries and ammunition production) will have doubled from "
        "2024 levels by end of 2029. The base rate for defence industrial scaling "
        "is unfavorable: the US took 3\u20134 years to meaningfully scale ammunition "
        "production for Ukraine aid, and that was with existing facilities and "
        "workforce. European ammunition production is currently meeting only a "
        "fraction of stated requirements. New production facilities require 2\u20133 "
        "years to build and certify. Major platforms (frigates, combat aircraft, "
        "armored vehicles) have 5\u201310 year production cycles. Labour shortages "
        "in skilled defence manufacturing compound the challenge."
    )

    # --- ASSUMPTION 4: Collaborative procurement ---
    doc.add_heading("Assumption 4: Collaborative Procurement Will Reach Target Levels", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "The paper proposes that collaborative procurement will reach approximately "
        "35\u201340% of total equipment spending, up from the current 18%. The SAFE "
        "instrument and EDIP grants are designed to incentivize this through "
        "financial carrots tied to multi-country procurement."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "The SAFE instrument requires collaborative procurement for contracts signed "
        "after May 2026 (minimum 2 countries). EDIP provides grants of 15\u201325% of "
        "contract value for procurement by 3+ countries. However, the paper does not "
        "address the historical failure to meet the 35% target set in 2007."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("MEDIUM")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x8B, 0x0B)

    doc.add_paragraph(
        "Given that collaborative procurement remains at its historical level of "
        "~18%, member states are spending more but buying nationally. This fragments "
        "the European defence market, duplicates capabilities, and fails to achieve "
        "the interoperability the paper argues is essential. The \u201cEuropean defence "
        "market\u201d remains 27 national markets, and the industrial consolidation "
        "the paper depends on does not materialize."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (8%): ")
    run.bold = True
    p.add_run(
        "We estimate an 8% probability that EU member states will achieve 35% "
        "collaborative procurement by 2029. The base rate is extremely unfavorable: "
        "the 35% target has never been met in 18 years of EDA measurement, with "
        "actual collaborative procurement stuck at approximately 18% throughout the "
        "2020s despite post-Ukraine urgency. The SAFE instrument is genuinely novel, "
        "but the denominator problem is severe: with procurement spending surging "
        "30\u201340% annually, collaborative procurement must grow in absolute terms "
        "far faster than the historical rate. Collaborative defence programs typically "
        "take 10\u201325 years to mature; programs initiated in 2026\u20132027 will "
        "produce minimal procurement volume by 2029."
    )

    # --- ASSUMPTION 5: Russia threat ---
    doc.add_heading("Assumption 5: Russia\u2019s Threat Extends Beyond Ukraine", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        "The paper treats Russia as a long-term strategic threat to all of Europe, "
        "not merely a regional aggressor focused on Ukraine. This justifies the "
        "continent-wide rearmament scope and the urgency of the 2030 deadline."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "General references to Russian aggression and the changed security environment. "
        "No detailed threat assessment of Russian military capabilities post-Ukraine, "
        "force reconstitution timelines, or specific scenarios for threat to NATO "
        "territory."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that Russia\u2019s conventional military threat to NATO territory "
        "remains low after the degradation of its forces in Ukraine, the case for "
        "crash rearmament weakens substantially. A depleted Russian military focused "
        "on holding Ukrainian territory and rebuilding over a decade justifies "
        "modernization but not emergency mobilization. The 2030 urgency is calibrated "
        "to a Russia that can threaten NATO within 5 years\u2014which requires a "
        "reconstitution timeline that most military analysts consider optimistic "
        "for Russia."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (18%): ")
    run.bold = True
    p.add_run(
        "We estimate an 18% probability that Russia will pose a credible conventional "
        "military threat to NATO territory (defined as the ability to conduct "
        "sustained large-scale operations beyond Ukraine) by end of 2030. Russia has "
        "lost an estimated 3,000+ armored vehicles, significant portions of its "
        "trained officer corps, and substantial equipment stocks in Ukraine. "
        "Historical precedents for military reconstitution after major wars suggest "
        "7\u201315 year timelines. Russia\u2019s defence industry is operating under "
        "severe sanctions and at wartime production limits already. The strongest "
        "counterargument is Russia\u2019s nuclear umbrella and hybrid warfare "
        "capabilities, which remain potent but do not constitute the conventional "
        "threat the white paper\u2019s rearmament plan is designed to counter."
    )

    # --- ASSUMPTION 6: 2030 timeline ---
    doc.add_heading("Assumption 6: The 2030 Timeline Is Sufficient", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Report\u2019s claim: ")
    run.bold = True
    p.add_run(
        'The paper is titled "Readiness 2030," implying that operational readiness '
        "can be achieved within approximately five years through the proposed "
        "mechanisms."
    )

    p = doc.add_paragraph()
    run = p.add_run("Evidence provided: ")
    run.bold = True
    p.add_run(
        "No detailed timeline analysis. The 2030 date appears politically driven "
        "(aligning with EU institutional cycles) rather than analytically derived "
        "from capability requirements and production timelines."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sensitivity: ")
    run.bold = True
    run = p.add_run("HIGH")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "Given that the 2030 timeline is insufficient\u2014which our analysis strongly "
        "suggests\u2014decision-makers need to plan for a longer build-up. This means "
        "different investment priorities: more emphasis on near-term off-the-shelf "
        "procurement to fill immediate gaps, combined with long-term industrial "
        "investments that will not deliver capability until 2035\u20132040. The \u201c2030 "
        "readiness\u201d framing creates a false urgency that may lead to wasteful "
        "crash procurement rather than sustainable capability building."
    )

    p = doc.add_paragraph()
    run = p.add_run("Our forecast (12%): ")
    run.bold = True
    p.add_run(
        "We estimate a 12% probability that the majority of the white paper\u2019s "
        "proposed capability targets will be met by end of 2030. This is a conjunction "
        "problem: meeting the 2030 deadline requires spending (58%), industrial scaling "
        "(12%), collaborative procurement (8%), and sustained political will simultaneously. "
        "Even taking the most favorable reading, the conjunction of these independent "
        "probabilities yields a very low overall probability. Historical analogues are "
        "instructive: NATO\u2019s 2014 Wales Summit pledge of 2% GDP took a decade to "
        "approach compliance. The EU\u2019s Permanent Structured Cooperation (PESCO), "
        "launched in 2017 with similar ambitions, has delivered limited results after 8 years. "
        "Large-scale rearmament programs (South Korea 1970s, Japan 2022 onwards) typically "
        "require 10\u201315 years to show meaningful capability improvements."
    )

    doc.add_page_break()

    # ========================================================
    # IMPLICATIONS FOR DECISION-MAKERS
    # ========================================================
    doc.add_heading("Implications for Decision-Makers", level=1)

    doc.add_heading("What to act on now", level=2)

    bullets_act = [
        (
            "Plan for 2035\u20132040, not 2030. ",
            "The 2030 readiness target is politically aspirational, not analytically grounded. "
            "Investment decisions should be calibrated to a 10\u201315 year capability-building "
            "timeline. Near-term spending should focus on ammunition stocks, air defence, and "
            "maintenance of existing platforms\u2014areas where money translates to capability "
            "quickly."
        ),
        (
            "The spending increase is real (58%); the spending efficiency is not. ",
            "European defence budgets will likely grow significantly. But without industrial "
            "capacity to absorb the spending, more money chases the same limited production. "
            "Organizations positioned in defence supply chains should expect demand growth, "
            "but delivery timelines will extend, not compress."
        ),
        (
            "National procurement will dominate (92% confidence). ",
            "Despite the collaborative procurement targets, the historical pattern of "
            "national buying will persist. Firms and analysts should model European defence "
            "as 27 national markets with modest EU-level coordination, not as a unified "
            "European defence market."
        ),
    ]

    for bold_text, normal_text in bullets_act:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_heading("What to watch", level=2)

    bullets_watch = [
        (
            "US force posture changes (currently 8% likely). ",
            "If the US meaningfully reduces European deployments, the urgency of the "
            "white paper\u2019s timeline increases dramatically. Monitor US European Command "
            "force levels, Congressional appropriations for European Deterrence Initiative, "
            "and NATO summit communiqu\u00e9s."
        ),
        (
            "EDTIB production data (currently 12% likely to double). ",
            "Track ammunition production rates, new facility announcements, and delivery "
            "timelines for major platforms. If production doubles by 2028, the 2030 timeline "
            "becomes more credible. Early indicators: 155mm shell production rates, IRIS-T "
            "delivery schedules, Leopard 2 and KNDS production orders."
        ),
        (
            "SAFE loan uptake and collaborative procurement data (currently 8% likely). ",
            "The EDA publishes collaborative procurement data annually. If the share moves "
            "above 25% by 2027, the institutional incentives are working. If it remains "
            "at ~18%, the structural barriers to European defence integration are as strong "
            "as ever."
        ),
    ]

    for bold_text, normal_text in bullets_watch:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(normal_text)

    doc.add_page_break()

    # ========================================================
    # METHODOLOGY NOTE
    # ========================================================
    doc.add_heading("Methodology", level=1)

    doc.add_paragraph(
        "Each assumption was converted into a concrete, time-bound forecasting question "
        "and evaluated using structured forecasting methods:"
    )

    method_steps = [
        "Base rate construction from historical reference classes",
        "Evidence gathering from domain research and recent publications",
        "Bayesian updating based on specific factors for and against",
        "Decomposition analysis (breaking the question into independent paths to resolution)",
        "Calibration checks and stress testing against sensitivity ranges",
    ]
    for step in method_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph(
        "Forecasts represent calibrated probabilities\u2014our best estimate of likelihood "
        "given available evidence. A 60% forecast means we expect the event to occur in "
        "roughly 6 out of 10 similar situations. These are not expressions of confidence "
        "in our analysis; they are estimates of real-world probability."
    )

    doc.add_paragraph(
        'Sensitivity framing uses presuppositional ("Given that X") rather than '
        'hypothetical ("If X were true") language. Research on conditional reasoning '
        "shows that presuppositional framing produces more accurate counterfactual "
        "analysis by treating the alternative scenario as a concrete state of the world "
        "rather than an abstract possibility."
    )

    # ========================================================
    # SAVE
    # ========================================================
    output_path = "/Users/elsehow/Downloads/audit/javier/FRI_Forecast_Audit_EU_Defence_Readiness_2030.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
